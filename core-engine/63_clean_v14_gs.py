import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score
import warnings
from scipy import stats
import os
import requests
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog
from datetime import datetime
from tenacity import retry, stop_after_attempt, wait_exponential
from dataclasses import dataclass, field
from enum import Enum, auto
from typing import List, Dict, Optional, Tuple, Any

warnings.filterwarnings("ignore", category=UserWarning)

class Severity(Enum):
    ERROR = auto()
    WARNING = auto()
    INFO = auto()

@dataclass
class ValidationIssue:
    severity: Severity
    category: str
    message: str
    affected_rows: Optional[List[int]] = None
    suggestion: Optional[str] = None

    def to_string(self) -> str:
        icon = "" if self.severity == Severity.ERROR else "" if self.severity == Severity.WARNING else ""
        lines = [f"   {icon} [{self.category}] {self.message}"]
        if self.affected_rows:
            rows_str = str(self.affected_rows[:10])
            if len(self.affected_rows) > 10:
                rows_str = rows_str[:-1] + f", ... 等共{len(self.affected_rows)}行]"
            lines.append(f"影响行索引: {rows_str}")
        if self.suggestion:
            lines.append(f"建议: {self.suggestion}")
        return "\n".join(lines)

@dataclass
class ValidationResult:
    passed: bool = True
    issues: List[ValidationIssue] = field(default_factory=list)
    quality_score: float = 100.0
    repair_log: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    def has_fatal_error(self) -> bool:
        return any(i.severity == Severity.ERROR for i in self.issues)

    def get_errors(self) -> List[ValidationIssue]:
        return [i for i in self.issues if i.severity == Severity.ERROR]

    def get_warnings(self) -> List[ValidationIssue]:
        return [i for i in self.issues if i.severity == Severity.WARNING]

    def get_infos(self) -> List[ValidationIssue]:
        return [i for i in self.issues if i.severity == Severity.INFO]

    def add_issue(self, issue: ValidationIssue):
        self.issues.append(issue)
        if issue.severity == Severity.ERROR:
            self.passed = False

    def print_sep(self, length=70):
        print(f"{'='*length}")

    def print_report(self, project_name: str):
        print()
        print(f"数据校验报告: {project_name}")
        score_color = "" if self.quality_score >= 80 else "" if self.quality_score >= 60 else ""
        print(f"{score_color} 数据质量评分: {self.quality_score:.1f}/100")

        if self.metadata:
            print(f"数据概览:")
            for key, value in self.metadata.items():
                print(f"   • {key}: {value}")

        errors = self.get_errors()
        warnings = self.get_warnings()
        infos = self.get_infos()

        if errors:
            print(f"错误 ({len(errors)}项)")
            for e in errors:
                print(e.to_string())

        if warnings:
            print(f"警告 ({len(warnings)}项) - 已默认填充:")
            for w in warnings:
                print(w.to_string())

        if infos:
            print(f"提示 ({len(infos)}项):")
            for i in infos:
                print(i.to_string())

        if self.repair_log:
            print(f"自动修复记录 ({len(self.repair_log)}项):")
            for log in self.repair_log:
                print(f"   • {log}")
        status = " 通过" if self.passed else " 失败"
        print(f"校验结果: {status}")
        print()

TREND_WEIGHT = 0.2
CONFIDENCE_LEVEL = 0.99
STRESS_STEP = 0.05
MAX_STRESS_FACTOR = 0.5
STRESS_SENSITIVITY = 0.8
MACRO_WEIGHTS = {
    'Housing_Volume': 0.2,
    'Housing_Price': 0.2,
    'Auction_Volume': 0.2,
    'Auction_Price': 0.2,
    'RE_Climate': 0.2
}

def load_batch_data():
    root = tk.Tk()
    root.withdraw()
    print("选择项目数据")
    proj_path = filedialog.askopenfilename(title="选择项目数据")
    if not proj_path:
        return None, None
    print("选择市场基准数据")
    mkt_path = filedialog.askopenfilename(title="选择市场基准数据")
    if not mkt_path:
        return None, None
    projects_dict = pd.read_excel(proj_path, sheet_name=None)
    market_df = pd.read_excel(mkt_path)
    return projects_dict, market_df

def preprocess_data(df, name_map, is_project_data=False):
    df_copy = df.copy()
    df_copy.rename(columns=name_map, inplace=True)
    df_copy['Date'] = pd.to_datetime(df_copy['Date'])
    df_copy['Total_Recovery'] = pd.to_numeric(df_copy['Total_Recovery'], errors='coerce').fillna(
        1).replace(0, 1)
    df_copy['Recovery_Rate'] = (df_copy['Actual_Recovery'] / df_copy['Total_Recovery']) * 100

    if 'Project_ID' in df_copy.columns:
        df_copy['Month_Since_Start'] = df_copy.groupby('Project_ID').cumcount()
    elif is_project_data:
        df_copy['Month_Since_Start'] = np.arange(len(df_copy))

    group_col = 'Project_ID' if 'Project_ID' in df_copy.columns else None
    if group_col:
        df_copy['Cum_Recovery'] = df_copy.groupby(group_col)['Actual_Recovery'].cumsum()
    else:
        df_copy['Cum_Recovery'] = df_copy['Actual_Recovery'].cumsum()

    df_copy['Cum_Recovery_Rate'] = (df_copy['Cum_Recovery'] / df_copy['Total_Recovery']) * 100

    macro_cols = ['Housing_Volume_Index', 'Housing_Price_Index', 'Auction_Volume_Index', 'Auction_Price_Index',
                  'RE_Climate_Index']
    for col in macro_cols:
        if col not in df_copy.columns:
            df_copy[col] = 100
        else:
            df_copy[col] = pd.to_numeric(df_copy[col], errors='coerce').fillna(100)

    return df_copy

def calculate_benchmarks_by_lifecycle(bank, market_df):
    def calc_avg(col, market_name, bank_name):
        market_avg = market_df.groupby('Month_Since_Start')[col].mean().reset_index().rename(
            columns={col: market_name})
        same_bank = market_df[market_df['Bank'] == bank]
        if not same_bank.empty:
            bank_avg = same_bank.groupby('Month_Since_Start')[col].mean().reset_index().rename(
                columns={col: bank_name})
        else:
            bank_avg = market_avg.rename(columns={market_name: bank_name})
        return market_avg, bank_avg

    market_p, bank_p = calc_avg('Recovery_Rate', 'Market_Average_Rate', 'Same_Bank_Peer_Rate')
    market_c, bank_c = calc_avg('Cum_Recovery_Rate', 'Market_Cumulative_Average_Rate',
                                  'Same_Bank_Cumulative_Peer_Rate')

    benchmarks = pd.merge(pd.merge(market_p, bank_p, on='Month_Since_Start', how='left'),
                          pd.merge(market_c, bank_c, on='Month_Since_Start', how='left'),
                          on='Month_Since_Start', how='left')
    benchmarks.ffill().bfill(inplace=True)
    return benchmarks

def forecast_future_hybrid(hist_df, market_df, bank, z_score, macro_stress_factor=1.0, periods=5,
                           trend_lookback=6):
    
    mkt_curve = market_df.groupby('Month_Since_Start')['Recovery_Rate'].mean()
    bank_peers = market_df[market_df['Bank'] == bank]
    same_bank_lifecycle = bank_peers.groupby('Month_Since_Start')[
        'Recovery_Rate'].mean() if not bank_peers.empty else mkt_curve

    last_period = hist_df['Month_Since_Start'].max()
    future_periods_abs = np.arange(last_period + 1, last_period + 1 + periods)

    # 预测期内的基准回收率
    future_benchmark_rates = pd.Series(
        0.7 * same_bank_lifecycle.reindex(future_periods_abs).values +
        0.3 * mkt_curve.reindex(future_periods_abs).values
    ).ffill().bfill().values

    historical_offset = (hist_df['Recovery_Rate'] - hist_df['Baseline_Rate']).tail(6).mean()
    lifecycle_forecast = future_benchmark_rates + historical_offset

    recent_data = hist_df.tail(trend_lookback)
    X_train = np.arange(len(recent_data)).reshape(-1, 1)
    y_train = recent_data['Recovery_Rate'].values

    if len(recent_data) < 3:
        r2_effective = 0
        trend_forecast = np.zeros(periods)
    else:
        model = LinearRegression()
        model.fit(X_train, y_train)
        y_pred_train = model.predict(X_train)
        r2 = r2_score(y_train, y_pred_train)
        r2_effective = max(0, r2)
        future_periods_rel = np.arange(len(recent_data), len(recent_data) + periods).reshape(-1, 1)
        trend_forecast = model.predict(future_periods_rel)

    dynamic_base_weight = TREND_WEIGHT * r2_effective
    decay_rate = 0.8
    period_weights = dynamic_base_weight * np.power(decay_rate, np.arange(1, periods + 1))
    final_forecast_periodic = period_weights * trend_forecast + (1 - period_weights) * lifecycle_forecast

    stress_adjustment = (1 - macro_stress_factor) * STRESS_SENSITIVITY
    final_forecast_periodic_stressed = final_forecast_periodic * (1 - stress_adjustment)
    final_forecast_periodic_stressed[final_forecast_periodic_stressed < 0] = 0

    last_actual_cumulative_rate = hist_df['Cum_Recovery_Rate'].iloc[-1]
    predicted_periodic_recovery = (final_forecast_periodic_stressed / 100) * \
                                  hist_df['Total_Recovery'].iloc[0]
    final_forecast_cumulative = last_actual_cumulative_rate + np.cumsum(
        predicted_periodic_recovery / hist_df['Total_Recovery'].iloc[0] * 100)

    periodic_volatility = hist_df['Deviation_Rate'].std()
    periodic_interval = z_score * periodic_volatility
    cumulative_volatility = hist_df['Cum_Deviation_Rate'].std()
    cumulative_interval = z_score * cumulative_volatility

    last_date = hist_df['Date'].max()
    future_dates = pd.to_datetime([last_date + pd.DateOffset(months=i * 3) for i in range(1, periods + 1)])

    future_market_cumulative = market_df.groupby('Month_Since_Start')['Cum_Recovery_Rate'].mean()
    future_same_bank_cumulative = bank_peers.groupby('Month_Since_Start')[
        'Cum_Recovery_Rate'].mean() if not bank_peers.empty else future_market_cumulative

    pred_cumulative_baseline = pd.Series(
        0.7 * future_same_bank_cumulative.reindex(future_periods_abs).values +
        0.3 * future_market_cumulative.reindex(future_periods_abs).values
    ).ffill().bfill().values

    last_market_multiplier = hist_df['Market_Multiplier'].iloc[
        -1] if 'Market_Multiplier' in hist_df else 1.0
    future_multiplier_stressed = 1 + ((last_market_multiplier - 1) * macro_stress_factor)

    pred_cumulative_lower_bound_stressed = pred_cumulative_baseline - z_score * cumulative_volatility * future_multiplier_stressed
    pred_cumulative_lower_bound_stressed[pred_cumulative_lower_bound_stressed < 0] = 0

    forecast_df = pd.DataFrame({
        'Date': future_dates,
        'Month_Since_Start': future_periods_abs,
        'Predicted_Rate': final_forecast_periodic_stressed,
        'Lower_Bound_Periodic': final_forecast_periodic_stressed - periodic_interval * future_multiplier_stressed,
        'Upper_Bound_Periodic': final_forecast_periodic_stressed + periodic_interval * future_multiplier_stressed,
        'Predicted_Cumulative_Rate': final_forecast_cumulative,
        'Lower_Bound_Cumulative': final_forecast_cumulative - cumulative_interval * future_multiplier_stressed,
        'Upper_Bound_Cumulative': final_forecast_cumulative + cumulative_interval * future_multiplier_stressed,
        'Predicted_Cumulative_Baseline': pred_cumulative_baseline,
        'Predicted_Cumulative_Lower_Bound_Stressed': pred_cumulative_lower_bound_stressed
    })

    forecast_df.loc[forecast_df['Lower_Bound_Periodic'] < 0, 'Lower_Bound_Periodic'] = 0
    forecast_df.loc[forecast_df['Lower_Bound_Cumulative'] < 0, 'Lower_Bound_Cumulative'] = 0

    return forecast_df

from scipy.optimize import minimize
from sklearn.linear_model import Ridge, ElasticNet
from sklearn.preprocessing import PolynomialFeatures
from sklearn.pipeline import Pipeline

class LifecycleForecaster:

    def __init__(self, market_df: pd.DataFrame, bank: str, same_bank_weight: float = 0.7):
        self.market_df = market_df
        self.bank = bank
        self.same_bank_weight = same_bank_weight
        self.historical_offset = 0
        self.residuals = []

    def get_bank_peers(self):
        return self.market_df[self.market_df['Bank'] == self.bank]

    def get_lifecycle_curve(self, bank_peers):
        mkt_curve = self.market_df.groupby('Month_Since_Start')['Recovery_Rate'].mean()
        if not bank_peers.empty:
            return bank_peers.groupby('Month_Since_Start')['Recovery_Rate'].mean(), mkt_curve
        return mkt_curve, mkt_curve

    def fit(self, hist_df):
        self.historical_offset = (hist_df['Recovery_Rate'] - hist_df['Baseline_Rate']).tail(6).mean()
        fitted = hist_df['Baseline_Rate'] + self.historical_offset
        self.residuals = (hist_df['Recovery_Rate'] - fitted).values
        return self

    def predict(self, periods) -> np.ndarray:
        bank_peers = self.get_bank_peers()
        same_bank_lifecycle, mkt_curve = self.get_lifecycle_curve(bank_peers)

        last_period = self.market_df['Month_Since_Start'].max() if 'Month_Since_Start' in self.market_df.columns else 0
        future_periods = np.arange(last_period + 1, last_period + 1 + periods)

        benchmark = (self.same_bank_weight * same_bank_lifecycle.reindex(future_periods).values +
                     (1 - self.same_bank_weight) * mkt_curve.reindex(future_periods).values)
        result = pd.Series(benchmark).ffill().bfill().values + float(self.historical_offset)
        return np.asarray(result).flatten()

    def get_confidence_intervals(self, periods, confidence) -> Tuple[np.ndarray, np.ndarray]:
        if len(self.residuals) < 5:
            std = np.std(self.residuals) if len(self.residuals) > 0 else 0.02
            z = stats.norm.ppf((1 + confidence) / 2)
            pred = self.predict(periods)
            return pred - z * std, pred + z * std

        np.random.seed(42)
        predictions = self.predict(periods)
        bootstrap_noise = np.random.choice(self.residuals, size=(1000, periods), replace=True)

        alpha = 1 - confidence
        lower = predictions + np.percentile(bootstrap_noise, alpha / 2 * 100, axis=0)
        upper = predictions + np.percentile(bootstrap_noise, (1 - alpha / 2) * 100, axis=0)
        return lower, upper

    @property
    def model_score(self) -> float:
        if len(self.residuals) == 0:
            return 0
        score = max(0, 1 - np.std(self.residuals) / 0.05)
        return score

class TrendForecaster:

    def __init__(self, lookback = None, poly_degree = 1,
                 regularization: str = 'ridge', alpha = 1.0):
        self.lookback = lookback
        self.poly_degree = min(poly_degree, 3)
        self.regularization = regularization
        self.alpha = alpha
        self.model = None
        self.r2 = 0
        self.residuals = []
        self.X_train = None
        self.y_train = None

    def get_optimal_lookback(self, hist_df) -> int:
        if self.lookback is not None:
            return min(self.lookback, len(hist_df))

        best_aic = float('inf')
        best_lookback = 6

        for lookback in [4, 6, 8, 10, len(hist_df)]:
            if lookback > len(hist_df) - 2:
                continue

            recent = hist_df.tail(lookback)
            X = np.arange(len(recent)).reshape(-1, 1)
            y = recent['Recovery_Rate'].values

            if self.poly_degree > 1:
                X = PolynomialFeatures(degree=self.poly_degree).fit_transform(X)

            try:
                if self.regularization == 'ridge':
                    model = Ridge(alpha=self.alpha)
                else:
                    model = ElasticNet(alpha=self.alpha, l1_ratio=0.5)
                model.fit(X, y)
                y_pred = model.predict(X)

                n = len(y)
                k = X.shape[1] + 1
                rss = np.sum((y - y_pred) ** 2)
                aic = n * np.log(rss / n + 1e-10) + 2 * k

                if aic < best_aic:
                    best_aic = aic
                    best_lookback = lookback
            except:
                continue

        return best_lookback

    def fit(self, hist_df):
        lookback = self.get_optimal_lookback(hist_df)
        recent = hist_df.tail(lookback)

        self.X_train = np.arange(len(recent)).reshape(-1, 1)
        self.y_train = recent['Recovery_Rate'].values

        steps = []
        if self.poly_degree > 1:
            steps.append(('poly', PolynomialFeatures(degree=self.poly_degree)))

        if self.regularization == 'ridge':
            steps.append(('model', Ridge(alpha=self.alpha)))
        elif self.regularization == 'elasticnet':
            steps.append(('model', ElasticNet(alpha=self.alpha, l1_ratio=0.5)))
        else:
            steps.append(('model', LinearRegression()))

        self.model = Pipeline(steps)
        self.model.fit(self.X_train, self.y_train)

        y_pred = self.model.predict(self.X_train)
        self.r2 = r2_score(self.y_train, y_pred)
        self.residuals = self.y_train - y_pred

        return self

    def predict(self, periods) -> np.ndarray:
        X_future = np.arange(len(self.X_train), len(self.X_train) + periods).reshape(-1, 1)
        pred = self.model.predict(X_future)
        return np.asarray(pred).flatten()[:periods]

    def get_confidence_intervals(self, periods, confidence) -> Tuple[np.ndarray, np.ndarray]:
        predictions = self.predict(periods)
        time_multiplier = np.sqrt(1 + np.arange(periods) / periods)
        std = np.std(self.residuals) if len(self.residuals) > 5 else 0.02
        z = stats.norm.ppf((1 + confidence) / 2)
        margin = z * std * time_multiplier
        return predictions - margin, predictions + margin

    @property
    def model_score(self) -> float:
        return max(0, self.r2)

class ExponentialSmoothingForecaster:

    def __init__(self, seasonal_periods = None):
        self.seasonal_periods = seasonal_periods
        self.alpha = 0.3
        self.level = None
        self.trend = None
        self.seasonal = None
        self.fitted = None
        self.residuals = []

    def fit(self, hist_df):
        y = hist_df['Recovery_Rate'].values
        n = len(y)

        self.level = y[0]
        self.trend = (y[1] - y[0]) if n > 1 else 0

        if self.seasonal_periods and n >= 2 * self.seasonal_periods:
            self.seasonal = np.zeros(self.seasonal_periods)
            for i in range(self.seasonal_periods):
                self.seasonal[i] = np.mean(y[i::self.seasonal_periods]) - np.mean(y)
        else:
            self.seasonal = None

        def objective(alpha):
            if alpha <= 0 or alpha >= 1:
                return float('inf')
            try:
                fitted = self.fit_values(y, alpha)
                return np.mean((y - fitted) ** 2)
            except:
                return float('inf')

        try:
            result = minimize(objective, x0=0.3, bounds=[(0.01, 0.99)], method='L-BFGS-B')
            self.alpha = result.x[0] if result.success else 0.3
        except:
            self.alpha = 0.3

        self.fitted = self.fit_values(y, self.alpha)
        self.residuals = y - self.fitted

        return self

    def fit_values(self, y, alpha) -> np.ndarray:
        n = len(y)
        fitted = np.zeros(n)
        level = float(self.level)
        trend = float(self.trend)

        for t in range(n):
            if self.seasonal is not None:
                s_idx = t % len(self.seasonal)
                seasonal = float(self.seasonal[s_idx])
            else:
                seasonal = 0

            fitted[t] = level + trend + seasonal

            observed = float(y[t])
            newlevel = alpha * (observed - seasonal) + (1 - alpha) * (level + trend)
            trend = 0.1 * (newlevel - level) + 0.9 * trend
            level = newlevel

            if self.seasonal is not None:
                self.seasonal[s_idx] = 0.1 * (observed - level) + 0.9 * seasonal

        self.level = level
        self.trend = trend
        return fitted

    def predict(self, periods) -> np.ndarray:
        predictions = []
        level = float(self.level)
        trend = float(self.trend)

        for h in range(1, periods + 1):
            if self.seasonal is not None:
                s_idx = (len(self.fitted) + h - 1) % len(self.seasonal)
                seasonal = float(self.seasonal[s_idx])
            else:
                seasonal = 0

            pred = level + h * trend + seasonal
            predictions.append(pred)

        result = np.array(predictions)
        return np.asarray(result).flatten()

    def get_confidence_intervals(self, periods, confidence) -> Tuple[np.ndarray, np.ndarray]:
        predictions = self.predict(periods)
        std = np.std(self.residuals) if len(self.residuals) > 0 else 0.02
        z = stats.norm.ppf((1 + confidence) / 2)
        margin = z * std * np.sqrt(np.arange(1, periods + 1))
        return predictions - margin, predictions + margin

    @property
    def model_score(self) -> float:
        if len(self.residuals) == 0:
            return 0
        mse = np.mean(self.residuals ** 2)
        return max(0, 1 - mse / 0.01)

class LogisticDecayForecaster:

    def __init__(self, target_recovery: float = 95.0):
        self.target_recovery = target_recovery
        self.params = None
        self.residuals = []
        self.fitted = None

    def logistic(self, t, L, k, t0) -> np.ndarray:
        return L / (1 + np.exp(-k * (t - t0)))

    def fit(self, hist_df):
        "拟合Logistic曲线到累计回收率"
        t = hist_df['Month_Since_Start'].values
        y = hist_df['Cum_Recovery_Rate'].values

        L0 = min(self.target_recovery, y[-1] * 1.5)
        k0 = 0.2
        t0 = len(y) * 0.6

        try:
            from scipy.optimize import curve_fit
            
            bounds = ([y[-1], 0.05, 0], [100.0, 2.0, len(y) * 3])
            popt, _ = curve_fit(self.logistic, t, y, p0=[L0, k0, t0],
                              bounds=bounds, maxfev=5000)
            self.params = popt
        except Exception:
            
            self.params = [L0, k0, t0]

        self.fitted = self.logistic(t, *self.params)
        self.residuals = y - self.fitted

        return self

    def predict(self, periods) -> np.ndarray:
        "预测未来各期的当期回收率"
        if self.params is None:
            return np.zeros(periods)

        L, k, t0 = self.params

        current_t = len(self.fitted) - 1 if self.fitted is not None else 0

        # 预测累计回收率
        future_t = np.arange(current_t + 1, current_t + 1 + periods)
        cumulative_future = self.logistic(future_t, L, k, t0)

        current_cumulative = self.fitted[-1] if self.fitted is not None else 0
        periodic_rates = np.diff(np.concatenate([[current_cumulative], cumulative_future]))

        return np.maximum(periodic_rates, 0).flatten()

    def predict_cumulative(self, periods) -> np.ndarray:
        "预测未来各期的累计回收率"
        if self.params is None:
            return np.zeros(periods)

        L, k, t0 = self.params
        current_t = len(self.fitted) - 1 if self.fitted is not None else 0
        future_t = np.arange(current_t + 1, current_t + 1 + periods)
        return self.logistic(future_t, L, k, t0).flatten()

    def get_confidence_intervals(self, periods, confidence) -> Tuple[np.ndarray, np.ndarray]:
        "基于Bootstrap残差计算置信区间"
        predictions = self.predict(periods)

        if len(self.residuals) < 5:
            std = np.std(self.residuals) if len(self.residuals) > 0 else 0.02
            z = stats.norm.ppf((1 + confidence) / 2)
            return predictions - z * std, predictions + z * std

        # Bootstrap
        np.random.seed(42)
        bootstrap_preds = []
        for _ in range(1000):
            noise = np.random.choice(self.residuals, size=periods, replace=True)
            bootstrap_preds.append(predictions + noise)

        bootstrap_preds = np.array(bootstrap_preds)
        alpha = 1 - confidence
        lower = np.percentile(bootstrap_preds, alpha / 2 * 100, axis=0)
        upper = np.percentile(bootstrap_preds, (1 - alpha / 2) * 100, axis=0)

        return lower, upper

    @property
    def model_score(self) -> float:
        if len(self.residuals) == 0 or self.fitted is None:
            return 0
        y = self.fitted + self.residuals
        ss_res = np.sum(self.residuals ** 2)
        ss_tot = np.sum((y - np.mean(y)) ** 2)
        r2 = 1 - ss_res / (ss_tot + 1e-10)
        return max(0, r2)

    @property
    def decay_rate(self) -> float:
        return self.params[1] if self.params is not None else 0.3

class BayesianHierarchicalForecaster:

    def __init__(self, market_df: pd.DataFrame, bank: str,
                 same_bank_weight: float = 0.7, prior_strength: float = 5.0):
        self.market_df = market_df
        self.bank = bank
        self.same_bank_weight = same_bank_weight
        self.prior_strength = prior_strength
        self.posterior_offset = 0
        self.residuals = []
        self.effective_weight = same_bank_weight

    def get_same_bank_or_market_avg(self, column=None):
        "获取同银行平均值，如不存在则返回市场平均"
        bank_peers = self.market_df[self.market_df['Bank'] == self.bank]
        if not bank_peers.empty:
            return bank_peers.groupby('Month_Since_Start')[column].mean()
        return self.market_df.groupby('Month_Since_Start')[column].mean()

    def fit(self, hist_df):
        "通过贝叶斯更新计算后验"
        n_obs = len(hist_df)

        prior_rates = self.get_same_bank_or_market_avg('Recovery_Rate')

        # 对齐到当前项目的期数
        periods = hist_df['Month_Since_Start'].values
        prior_values = prior_rates.reindex(periods).values

        # 观测值
        obs_values = hist_df['Recovery_Rate'].values

        prior_std = prior_rates.std() if len(prior_rates) > 1 else 0.05
        obs_std = hist_df['Recovery_Rate'].std() if n_obs > 1 else 0.05

        # 贝叶斯更新：精度加权
        prior_precision = 1 / (prior_std ** 2 + 1e-6)
        obs_precision = n_obs / (obs_std ** 2 + 1e-6)

        # 后验均值
        deviations = obs_values - prior_values
        mean_deviation = np.mean(deviations)

        # 加权平均：先验 vs 观测
        total_precision = self.prior_strength * prior_precision + obs_precision
        posterior_deviation = (self.prior_strength * prior_precision * 0 +
                               obs_precision * mean_deviation) / total_precision

        self.posterior_offset = posterior_deviation

        self.effective_weight = obs_precision / total_precision

        fitted = prior_values + posterior_deviation
        self.residuals = obs_values - fitted

        return self

    def predict(self, periods) -> np.ndarray:
        "预测未来各期"
        same_bank_avg = self.get_same_bank_or_market_avg('Recovery_Rate')
        market_avg = self.market_df.groupby('Month_Since_Start')['Recovery_Rate'].mean()
        last_period = self.market_df['Month_Since_Start'].max() if 'Month_Since_Start' in self.market_df.columns else 0
        current_period = len(self.residuals) if self.residuals is not None else 0
        future_periods = np.arange(current_period, current_period + periods)
        baseline = (self.same_bank_weight * same_bank_avg.reindex(future_periods).values +
                    (1 - self.same_bank_weight) * market_avg.reindex(future_periods).values)
        baseline = pd.Series(baseline).ffill().bfill().values
        predictions = baseline + self.posterior_offset
        return np.maximum(predictions, 0).flatten()

    def get_confidence_intervals(self, periods, confidence) -> Tuple[np.ndarray, np.ndarray]:
        "计算置信区间"
        predictions = self.predict(periods)
        std = np.std(self.residuals) if len(self.residuals) > 0 else 0.02
        time_multiplier = np.sqrt(1 + np.arange(periods) / periods)
        z = stats.norm.ppf((1 + confidence) / 2)
        margin = z * std * time_multiplier

        return predictions - margin, predictions + margin

    @property
    def model_score(self) -> float:
        if len(self.residuals) == 0:
            return 0
        # 有效样本越多、残差越小，分数越高
        sample_bonus = min(0.3, len(self.residuals) / 30)
        fit_quality = max(0, 1 - np.std(self.residuals) / 0.05)
        return 0.5 * fit_quality + 0.5 * sample_bonus

    @property
    def bayesian_weight(self) -> float:
        return self.effective_weight

class StageStratifiedForecaster:

    STAGE_EARLY = 'early'
    STAGE_MIDDLE = 'middle'
    STAGE_LATE = 'late'

    # 累计回收率阈值
    RECOVERY_THRESHOLDS = [0.25, 0.60]

    def __init__(self, market_df: pd.DataFrame, bank: str):
        self.market_df = market_df
        self.bank = bank
        self.curr_stage = None
        self.stage_models = {}
        self.residuals = []
        
        self.STAGE_PERIODS = self.calculate_stage_thresholds()
        print(f" 数据驱动阶段划分: 早期(0-{self.STAGE_PERIODS[self.STAGE_EARLY][1]}), "
              f"中期({self.STAGE_PERIODS[self.STAGE_MIDDLE][0]}-{self.STAGE_PERIODS[self.STAGE_MIDDLE][1]}), "
              f"晚期({self.STAGE_PERIODS[self.STAGE_LATE][0]}+)")

    def calculate_stage_thresholds(self) -> dict:
        """
        数据驱动确定阶段临界点
        基于同银行历史项目的中位生存时间（达到特定累计回收率所需期数）
        自动计算阶段转换点 τ_k
        """
        same_bank_projects = self.market_df[self.market_df['Bank'] == self.bank]

        if same_bank_projects.empty or len(same_bank_projects['Project_ID'].unique()) < 3:
            # 历史数据不足，使用保守默认值
            return {
                self.STAGE_EARLY: (0, 4),
                self.STAGE_MIDDLE: (5, 12),
                self.STAGE_LATE: (13, float('inf'))
            }

        survival_times = {threshold: [] for threshold in self.RECOVERY_THRESHOLDS}
        for project_id in same_bank_projects['Project_ID'].unique():
            project_data = same_bank_projects[same_bank_projects['Project_ID'] == project_id]
            project_data = project_data.sort_values('Month_Since_Start')

            if 'Cum_Recovery_Rate' not in project_data.columns:
                continue

            cumulative = project_data['Cum_Recovery_Rate'].values

            for threshold in self.RECOVERY_THRESHOLDS:
                exceed_indices = np.where(cumulative >= threshold * 100)[0]
                if len(exceed_indices) > 0:
                    survival_time = project_data.iloc[exceed_indices[0]]['Month_Since_Start']
                    survival_times[threshold].append(int(survival_time))

        # 计算中位生存时间作为临界点
        if survival_times[self.RECOVERY_THRESHOLDS[0]]:
            tau_1 = int(np.median(survival_times[self.RECOVERY_THRESHOLDS[0]]))
        else:
            tau_1 = 4
        if survival_times[self.RECOVERY_THRESHOLDS[1]]:
            tau_2 = int(np.median(survival_times[self.RECOVERY_THRESHOLDS[1]]))
        else:
            tau_2 = 12
        return {
            self.STAGE_EARLY: (0, tau_1),
            self.STAGE_MIDDLE: (tau_1 + 1, tau_2),
            self.STAGE_LATE: (tau_2 + 1, float('inf'))
        }

    def identify_stage(self, period: int) -> str:
        "识别当前阶段"
        for stage, (start, end) in self.STAGE_PERIODS.items():
            if start <= period <= end:
                return stage
        return self.STAGE_LATE

    def fit_stage_model(self, stage: str, hist_df):
        stage_range = self.STAGE_PERIODS[stage]

        # 从市场数据中提取该阶段的历史表现
        bank_peers = self.market_df[self.market_df['Bank'] == self.bank]

        if not bank_peers.empty:
            stage_data = bank_peers[
                (bank_peers['Month_Since_Start'] >= stage_range[0]) &
                (bank_peers['Month_Since_Start'] <= stage_range[1])
            ]
        else:
            stage_data = self.market_df[
                (self.market_df['Month_Since_Start'] >= stage_range[0]) &
                (self.market_df['Month_Since_Start'] <= stage_range[1])
            ]

        if len(stage_data) == 0:
            return {'mean': 0.02, 'std': 0.01}

        periodic_rates = stage_data.groupby('Month_Since_Start')['Recovery_Rate'].mean()

        return {
            'mean': periodic_rates.mean(),
            'std': periodic_rates.std() if len(periodic_rates) > 1 else 0.01,
            'trend': 0
        }

    def fit(self, hist_df):
        current_period = hist_df['Month_Since_Start'].iloc[-1]
        self.curr_stage = self.identify_stage(current_period)

        for stage in [self.STAGE_EARLY, self.STAGE_MIDDLE, self.STAGE_LATE]:
            self.stage_models[stage] = self.fit_stage_model(stage, hist_df)

        fitted = []
        for _, row in hist_df.iterrows():
            stage = self.identify_stage(row['Month_Since_Start'])
            fitted.append(self.stage_models[stage]['mean'])

        self.residuals = hist_df['Recovery_Rate'].values - np.array(fitted)

        return self

    def predict(self, periods) -> np.ndarray:
        "预测未来各期"
        predictions = []
        current_period = len(self.residuals) if self.residuals is not None else 0

        for i in range(periods):
            future_period = current_period + i
            stage = self.identify_stage(future_period)
            model = self.stage_models.get(stage, {'mean': 0.01, 'std': 0.01})
            base_rate = model['mean']
            # 晚期阶段：衰减调整
            if stage == self.STAGE_LATE:
                late_periods = max(0, future_period - 15)  # 从15期开始计算
                base_rate *= (0.96 ** late_periods)

            predictions.append(base_rate)

        if len(self.residuals) > 0:
            avg_residual = np.mean(self.residuals[-6:])  # 最近6期平均
            predictions = [max(0, p + avg_residual) for p in predictions]

        return np.array(predictions).flatten()

    def get_confidence_intervals(self, periods, confidence) -> Tuple[np.ndarray, np.ndarray]:
        "置信区间"
        predictions = self.predict(periods)
        z = stats.norm.ppf((1 + confidence) / 2)

        # 当前期数
        current_period = len(self.residuals) if self.residuals is not None else 0

        margins = []
        for i in range(periods):
            future_period = current_period + i
            stage = self.identify_stage(future_period)
            model = self.stage_models.get(stage, {'std': 0.02})
            time_factor = 1 + (future_period - current_period) / periods * 0.3
            margin = z * model['std'] * time_factor
            margins.append(margin)

        lower = predictions - np.array(margins)
        upper = predictions + np.array(margins)

        return np.maximum(lower, 0), upper

    @property
    def model_score(self) -> float:
        if len(self.residuals) == 0:
            return 0
        return max(0, 1 - np.std(self.residuals) / 0.05)

    @property
    def current_stage(self) -> str:
        return self.curr_stage

class ForecastCombiner:

    def __init__(self, forecasters: List, ensemble_method: str = 'adaptive_weight'):
        self.forecasters = forecasters
        self.ensemble_method = ensemble_method
        self.weights = None
        self.fitted_forecasters = []

    def fit(self, hist_df) -> 'ForecastCombiner':
        self.fitted_forecasters = []
        scores = []

        for forecaster in self.forecasters:
            try:
                fitted = forecaster.fit(hist_df)
                self.fitted_forecasters.append(fitted)
                scores.append(fitted.model_score)
            except Exception as e:
                print(f"预测器 {forecaster.__class__.__name__} 拟合失败: {e}")
                scores.append(0)

        scores = np.array(scores)

        if self.ensemble_method == 'equal':
            self.weights = np.ones(len(scores)) / len(scores)
        elif self.ensemble_method == 'best_only':
            best_idx = np.argmax(scores)
            self.weights = np.zeros(len(scores))
            self.weights[best_idx] = 1.0
        elif self.ensemble_method == 'adaptive_weight':
            if np.sum(scores) > 0:
                exp_scores = np.exp(scores - np.max(scores))
                self.weights = exp_scores / np.sum(exp_scores)
            else:
                self.weights = np.ones(len(scores)) / len(scores)
        elif self.ensemble_method == 'stacking':
            self.weights = self.compute_stacking_weights(hist_df)

        print(f"集成预测器权重分配")
        for f, w in zip(self.fitted_forecasters, self.weights):
            if w > 0.01:
                print(f"   {f.__class__.__name__}: {w:.2%} (评分: {f.model_score:.3f})")

        return self

    def compute_stacking_weights(self, hist_df) -> np.ndarray:
        fitted_values = []
        for forecaster in self.fitted_forecasters:
            fitted = forecaster.predict(len(hist_df))
            fitted_values.append(fitted)

        X = np.column_stack(fitted_values)
        y = hist_df['Recovery_Rate'].values

        def objective(weights):
            pred = X @ weights
            return np.mean((y - pred) ** 2)

        constraints = {'type': 'eq', 'fun': lambda w: np.sum(w) - 1}
        bounds = [(0, 1) for _ in range(len(self.fitted_forecasters))]

        result = minimize(objective, x0=np.ones(len(self.fitted_forecasters)) / len(self.fitted_forecasters),
                         bounds=bounds, constraints=constraints, method='SLSQP')

        return result.x if result.success else np.ones(len(self.fitted_forecasters)) / len(self.fitted_forecasters)

    def predict(self, periods) -> np.ndarray:
        predictions = np.zeros(periods)
        for forecaster, weight in zip(self.fitted_forecasters, self.weights):
            if weight > 0.001:
                pred = forecaster.predict(periods)
                pred = np.asarray(pred).flatten()[:periods]
                if len(pred) == periods:
                    predictions += weight * pred
        return predictions
    def get_confidence_intervals(self, periods, confidence) -> Tuple[np.ndarray, np.ndarray]:
        all_predictions = []
        for forecaster in self.fitted_forecasters:
            pred = forecaster.predict(periods)
            all_predictions.append(np.asarray(pred).flatten()[:periods])

        all_predictions = np.array(all_predictions)
        mean_pred = np.average(all_predictions, axis=0, weights=self.weights)

        variance = np.average((all_predictions - mean_pred.reshape(1, -1)) ** 2, axis=0, weights=self.weights)

        for i, forecaster in enumerate(self.fitted_forecasters):
            lower, upper = forecaster.get_confidence_intervals(periods, confidence)
            model_variance = ((upper - lower) / (2 * stats.norm.ppf((1 + confidence) / 2))) ** 2
            variance += self.weights[i] * model_variance

        z = stats.norm.ppf((1 + confidence) / 2)
        margin = z * np.sqrt(variance)
        return mean_pred - margin, mean_pred + margin

    def getmodel_diagnostics(self) -> Dict:
        return {
            'forecasters': [f.__class__.__name__ for f in self.fitted_forecasters],
            'weights': self.weights.tolist(),
            'scores': [f.model_score for f in self.fitted_forecasters]
        }

def forecast_future_enhanced(
    hist_df,
    market_df: pd.DataFrame,
    bank: str,
    z_score: float,
    macro_stress_factor: float = 1.0,
    periods = 5
) -> pd.DataFrame:

    print(f"启动预测模型 (预测期数: {periods}) ")

    forecasters = {
        'logistic': LogisticDecayForecaster(target_recovery=85.0),
        'bayesian': BayesianHierarchicalForecaster(market_df, bank,
                                                   same_bank_weight=0.7,
                                                   prior_strength=5.0),
        'stage': StageStratifiedForecaster(market_df, bank)
    }

    # 拟合各模型
    predictions = {}
    scores = {}

    for name, forecaster in forecasters.items():
        try:
            forecaster.fit(hist_df)
            predictions[name] = forecaster.predict(periods)
            scores[name] = forecaster.model_score
            print(f"{name:10s} (R²={scores[name]:.3f})")
        except Exception as e:
            print(f"{name:10s} 模型拟合失败:{e}")
            scores[name] = 0
            predictions[name] = np.zeros(periods)

    # 固定权重分配：核心模型
    weight_dict = {
        'logistic': 0.20,
        'bayesian': 0.25,
        'stage': 0.55
    }
    print(f"模型权重: Logistic={weight_dict['logistic']:.2f}, "
          f"Bayesian={weight_dict['bayesian']:.2f}, Stage={weight_dict['stage']:.2f}")

    # 计算加权预测
    base_forecast = np.zeros(periods)
    for name, pred in predictions.items():
        base_forecast += weight_dict[name] * pred

    base_forecast = np.maximum(base_forecast, 0)

    stress_adjustment = (1 - macro_stress_factor) * STRESS_SENSITIVITY
    stressed_forecast = base_forecast * (1 - stress_adjustment)
    stressed_forecast = np.maximum(stressed_forecast, 0)

    # 使用贝叶斯模型的置信区间
    bayesian_ci = forecasters['bayesian'].get_confidence_intervals(periods, CONFIDENCE_LEVEL)
    pred_matrix = np.array([predictions[name] for name in forecasters.keys()])
    model_disagreement = np.std(pred_matrix, axis=0)
    z = stats.norm.ppf((1 + CONFIDENCE_LEVEL) / 2)
    disagreement_margin = z * model_disagreement
    lower_periodic = np.maximum(stressed_forecast - disagreement_margin, 0)
    upper_periodic = stressed_forecast + disagreement_margin
    last_actual_cumulative = hist_df['Cum_Recovery_Rate'].iloc[-1]
    total_expected = hist_df['Total_Recovery'].iloc[0]
    predicted_periodic_recovery = (stressed_forecast / 100) * total_expected
    cumulative_forecast = last_actual_cumulative + np.cumsum(
        predicted_periodic_recovery / total_expected * 100
    )
    cumulative_variance = np.zeros(periods)
    for i in range(periods):
        periodic_var = ((upper_periodic[i] - lower_periodic[i]) / (2 * z_score)) ** 2
        cumulative_variance[i] = cumulative_variance[i-1] + periodic_var if i > 0 else periodic_var
    cumulative_margin = z_score * np.sqrt(cumulative_variance)
    cumulative_lower = np.maximum(cumulative_forecast - cumulative_margin, 0)
    cumulative_upper = cumulative_forecast + cumulative_margin
    current_period = hist_df['Month_Since_Start'].iloc[-1]
    current_stage = forecasters['stage'].current_stage
    stage_names = {'early': '早期', 'middle': '中期', 'late': '晚期'}
    last_date = hist_df['Date'].max()
    future_dates = pd.to_datetime([last_date + pd.DateOffset(months=i*3) for i in range(1, periods+1)])
    future_periods_abs = np.arange(current_period + 1, current_period + 1 + periods)
    forecast_df = pd.DataFrame({
        'Date': future_dates,
        'Month_Since_Start': future_periods_abs,
        'Predicted_Rate': stressed_forecast,
        'Lower_Bound_Periodic': lower_periodic,
        'Upper_Bound_Periodic': upper_periodic,
        'Predicted_Cumulative_Rate': cumulative_forecast,
        'Lower_Bound_Cumulative': cumulative_lower,
        'Upper_Bound_Cumulative': cumulative_upper
    })

    print(f" 当前所处阶段: {stage_names.get(current_stage, current_stage)}")
    print(f" Logistic衰减速率: {forecasters['logistic'].decay_rate:.3f}")
    print(f" 贝叶斯有效权重: {forecasters['bayesian'].bayesian_weight:.2f}")
    print(f" 预测完成，集成{len([s for s in scores.values() if s > 0])}个有效模型")
    return forecast_df

@dataclass
class StressScenario:
    name: str
    description: str
    housing_price_impact: float
    judicial_efficiency_impact: float
    macro_economic_impact: float
    severity: str

class StressTester:
    SCENARIOS = {
        # 房地产市场情景
        'housing_mild': StressScenario(
            '房价下跌10%', '房地产市场轻度调整，房价下跌10%',
            housing_price_impact=0.90, judicial_efficiency_impact=1.0, macro_economic_impact=0.95,
            severity='轻度'
        ),
        'housing_moderate': StressScenario(
            '房价下跌25%', '房地产市场中度下行，房价下跌25%',
            housing_price_impact=0.75, judicial_efficiency_impact=1.0, macro_economic_impact=0.90,
            severity='中度'
        ),
        'housing_severe': StressScenario(
            '房价下跌40%', '房地产市场重度危机，房价下跌40%',
            housing_price_impact=0.60, judicial_efficiency_impact=0.95, macro_economic_impact=0.80,
            severity='重度'
        ),

        # 司法处置效率情景
        'judicial_mild': StressScenario(
            '处置周期+3月', '司法处置轻度放缓，平均周期延长3个月',
            housing_price_impact=1.0, judicial_efficiency_impact=0.90, macro_economic_impact=1.0,
            severity='轻度'
        ),
        'judicial_moderate': StressScenario(
            '处置周期+6月', '司法处置中度放缓，平均周期延长6个月',
            housing_price_impact=1.0, judicial_efficiency_impact=0.80, macro_economic_impact=0.95,
            severity='中度'
        ),
        'judicial_severe': StressScenario(
            '处置停滞', '司法处置重度受阻，流拍率大幅上升',
            housing_price_impact=0.95, judicial_efficiency_impact=0.65, macro_economic_impact=0.90,
            severity='重度'
        ),

        # 宏观经济情景
        'macro_mild': StressScenario(
            '经济轻度衰退', '失业率上升2%，GDP增速放缓',
            housing_price_impact=0.95, judicial_efficiency_impact=1.0, macro_economic_impact=0.90,
            severity='轻度'
        ),
        'macro_moderate': StressScenario(
            '经济中度衰退', '失业率上升5%，GDP下滑3%',
            housing_price_impact=0.90, judicial_efficiency_impact=0.95, macro_economic_impact=0.75,
            severity='中度'
        ),
        'macro_severe': StressScenario(
            '系统性金融危机', '全面的信贷紧缩和经济危机',
            housing_price_impact=0.70, judicial_efficiency_impact=0.85, macro_economic_impact=0.60,
            severity='重度'
        ),

        # 组合压力情景
        'combined_moderate': StressScenario(
            '中度组合压力', '房价跌20% + 司法放缓 + 经济下行',
            housing_price_impact=0.80, judicial_efficiency_impact=0.85, macro_economic_impact=0.80,
            severity='中度'
        ),
        'combined_severe': StressScenario(
            '重度组合压力', '房价跌35% + 司法停滞 + 金融危机',
            housing_price_impact=0.65, judicial_efficiency_impact=0.60, macro_economic_impact=0.55,
            severity='重度'
        ),
    }

    def __init__(self, hist_df, market_df: pd.DataFrame,
                 bank: str, z_score: float, current_rating: str):
        self.hist_df = hist_df
        self.market_df = market_df
        self.bank = bank
        self.z_score = z_score
        self.current_rating = current_rating
        self.rating_map = {"正常类": 0, "关注类": 1, "次级类": 2, "可疑类": 3}
        self.thresholds = {"关注类": 1, "次级类": 5, "可疑类": 10}
        self.ratinglevels = ["正常类", "关注类", "次级类", "可疑类"]

    def print_sep(self, length=70):
        print(f"{'='*length}")

    def calculate_composite_stress_factor(self, scenario: StressScenario) -> float:
        """
        计算综合压力因子
        根据项目所处阶段，不同维度的压力影响权重不同：
        - 早期: 宏观因素更重要 (处置周期长)
        - 中期: 司法效率更重要 (大量进入司法程序)
        - 晚期: 房价因素更重要 (处置抵押物)
        """
        current_period = self.hist_df['Month_Since_Start'].iloc[-1]

        if current_period <= 4:  # 早期
            weights = {'housing': 0.3, 'judicial': 0.3, 'macro': 0.4}
        elif current_period <= 12:  # 中期
            weights = {'housing': 0.4, 'judicial': 0.4, 'macro': 0.2}
        else:
            weights = {'housing': 0.5, 'judicial': 0.3, 'macro': 0.2}

        composite = (
            weights['housing'] * scenario.housing_price_impact +
            weights['judicial'] * scenario.judicial_efficiency_impact +
            weights['macro'] * scenario.macro_economic_impact
        )

        return composite

    def run_scenario(self, scenario_key: str) -> Dict:
        """运行单个压力情景"""
        scenario = self.SCENARIOS[scenario_key]
        stress_factor = self.calculate_composite_stress_factor(scenario)
        forecast_df = forecast_future_enhanced(
            self.hist_df, self.market_df, self.bank, self.z_score,
            macro_stress_factor=stress_factor
        )

        final_pred = forecast_df.iloc[-1]
        return {
            'scenario': scenario,
            'stress_factor': stress_factor,
            'forecast': forecast_df,
            'final_cumulative_rate': final_pred['Predicted_Cumulative_Rate'],
            'final_lower_bound': final_pred['Lower_Bound_Cumulative'],
            'deviation': final_pred['Lower_Bound_Cumulative'] - final_pred['Predicted_Cumulative_Rate']
        }

    def find_critical_scenarios(self) -> Dict[str, Any]:
        "找出导致评级下调的临界情景"
        currentlevel = self.rating_map.get(self.current_rating, 0)
        critical_results = {}

        # 首先运行基准预测（无压力）
        baseline_result = self.run_scenario('housing_mild')
        baseline_cumulative = baseline_result['final_cumulative_rate']

        print()
        self.print_sep()
        print(f" NPL ABS多维度压力测试报告")
        self.print_sep()
        print(f"当前评级: {self.current_rating}")
        print(f"当前期数: 第{self.hist_df['Month_Since_Start'].iloc[-1]}期")
        print(f"基准累计回收率: {baseline_cumulative:.2f}%")

        # 按维度分组测试
        dimensions = {
            ' 房地产市场压力': ['housing_mild', 'housing_moderate', 'housing_severe'],
            ' 司法处置压力': ['judicial_mild', 'judicial_moderate', 'judicial_severe'],
            ' 宏观经济压力': ['macro_mild', 'macro_moderate', 'macro_severe'],
            ' 组合压力': ['combined_moderate', 'combined_severe']
        }

        for dim_name, scenario_keys in dimensions.items():
            print(f"\n{dim_name}")
            for key in scenario_keys:
                result = self.run_scenario(key)
                scenario = result['scenario']
                cumulative_decline = baseline_cumulative - result['final_cumulative_rate']
                downgrade_triggered = False
                for target_rating, threshold in self.thresholds.items():
                    targetlevel = self.rating_map[target_rating]
                    if targetlevel > currentlevel and cumulative_decline > threshold:
                        downgrade_triggered = True
                        if target_rating not in critical_results:
                            critical_results[target_rating] = []
                        critical_results[target_rating].append({
                            'scenario': scenario,
                            'stress_factor': result['stress_factor'],
                            'final_rate': result['final_cumulative_rate'],
                            'decline': cumulative_decline
                        })
                        break

                status = " 触发下调" if downgrade_triggered else " 安全"
                print(f"  {scenario.severity:2s} | {scenario.name:12s} | "
                      f"综合因子={result['stress_factor']:.2f} | "
                      f"预测累计={result['final_cumulative_rate']:.2f}% | "
                      f"下降{cumulative_decline:.2f}% | {status}")

        return critical_results
    def run_full_stress_test(self) -> Dict[str, float]:
        critical_scenarios = self.find_critical_scenarios()
        stress_results = {}
        for rating, scenarios in critical_scenarios.items():
            if scenarios:
                lightest = max(scenarios, key=lambda x: x['stress_factor'])
                stress_results[rating] = lightest['stress_factor']
        print()
        print("压力测试结论")
        if stress_results:
            print("以下情景将导致评级下调：")
            for rating, scenarios in critical_scenarios.items():
                if scenarios:
                    lightest = max(scenarios, key=lambda x: x['stress_factor'])
                    print(f"  - 下调至 '{rating}': 综合压力因子 {lightest['stress_factor']:.2f}, "
                          f"累计回收率下降 {lightest['decline']:.2f}%")
        else:
            print("所有压力情景下评级保持稳定")
        print()
        return stress_results

def run_stress_test(hist_df, market_df, bank, z_score, current_rating):
    """
    压力测试入口函数
    使用新的多维度压力测试引擎替代原有单一因子测试
    """
    engine = StressTester(hist_df, market_df, bank, z_score, current_rating)
    return engine.run_full_stress_test()

def plot_analysis_and_forecast(hist_df, forecast_df, project_name, confidencelevel,
                               risk_metrics=None, stress_results=None):
    """
    Goldman Sachs Research Style — 4-panel chart
    - 图1: 当期回收率监控与预测
    - 图2: 累计回收率偏离度分析
    - 图3: 压力测试结果热力图
    - 图4: 回收进度仪表盘
    """
    # === Goldman Sachs Color Palette ===
    NAVY = '#003A70'
    DARK_GOLD = '#C4A43E'
    STEEL_BLUE = '#5B9BD5'
    DARK_TEAL = '#2E75B6'
    DARK_RED = '#C0392B'
    DARK_GREEN = '#1E8449'
    SLATE = '#708090'
    DARK_TEXT = '#333333'
    MED_GRAY = '#999999'
    BORDER_GRAY = '#E0E0E0'

    # Global style
    plt.style.use('seaborn-v0_8-whitegrid')
    plt.rcParams.update({
        'figure.facecolor': 'white',
        'axes.facecolor': 'white',
        'axes.edgecolor': BORDER_GRAY,
        'axes.labelcolor': DARK_TEXT,
        'axes.titlecolor': NAVY,
        'text.color': DARK_TEXT,
        'xtick.color': MED_GRAY,
        'ytick.color': MED_GRAY,
        'grid.color': BORDER_GRAY,
        'grid.alpha': 0.5,
        'grid.linewidth': 0.4,
        'axes.spines.top': False,
        'axes.spines.right': False,
        'axes.linewidth': 0.6,
    })

    try:
        plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei']
        plt.rcParams['axes.unicode_minus'] = False
    except Exception as e:
        print(f"Mpl font warning: {e}")

    fig = plt.figure(figsize=(20, 24), facecolor='white')
    gs = fig.add_gridspec(3, 2, hspace=0.30, wspace=0.22,
                          height_ratios=[1, 1, 1], width_ratios=[1, 1])

    ax1 = fig.add_subplot(gs[0, 0])  # Period Recovery
    ax2 = fig.add_subplot(gs[0, 1])  # Cumulative Recovery
    ax5 = fig.add_subplot(gs[1, 0])  # Cashflow Waterfall
    ax6 = fig.add_subplot(gs[1, 1])  # Deviation Trend
    ax3 = fig.add_subplot(gs[2, 0])  # Stress Test Heatmap
    ax4 = fig.add_subplot(gs[2, 1])  # Recovery Gauge

    def period_formatter(x, pos):
        return f"Q{int(x) + 1}"

    def gs_style_ax(ax, title=None):
        """Apply GS style to an axes."""
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color(BORDER_GRAY)
        ax.spines['bottom'].set_color(BORDER_GRAY)
        ax.tick_params(colors=MED_GRAY, labelsize=9)
        ax.grid(True, alpha=0.4, linewidth=0.4)
        if title:
            ax.set_title(title, fontsize=13, fontweight='bold', color=NAVY, loc='left')

    x_hist = hist_df['Month_Since_Start']
    x_fore = forecast_df['Month_Since_Start']
    periods_hist = x_hist.values
    periods_fore = x_fore.values

    # ---- PANEL 1: Periodic Recovery Rate ----
    ax1.plot(x_hist, hist_df['Recovery_Rate'], label='Actual Recovery Rate (%)',
             color=NAVY, marker='o', zorder=5, linewidth=2, markersize=5)
    ax1.plot(x_hist, hist_df['Baseline_Rate'], label='Weighted Baseline',
             color=DARK_GOLD, linestyle='--', zorder=4, linewidth=1.5)
    ax1.fill_between(x_hist, hist_df['Lower_Bound_Rate'], hist_df['Upper_Bound_Rate'],
                     color=SLATE, alpha=0.12, label=f'Historical Range ({confidencelevel*100:.0f}% CI)')

    anomalies = hist_df[hist_df['Anomaly']]
    if not anomalies.empty:
        ax1.scatter(anomalies['Month_Since_Start'], anomalies['Recovery_Rate'],
                    color=DARK_RED, s=120, edgecolor='white', linewidths=0.5,
                    marker='X', label='Detected Anomalies', zorder=10)
        for idx, row in anomalies.iterrows():
            deviation = row['Recovery_Rate'] - row['Baseline_Rate']
            direction = "+" if deviation > 0 else ""
            ax1.annotate(f'{direction}{deviation:.1f}%',
                        xy=(row['Month_Since_Start'], row['Recovery_Rate']),
                        xytext=(5, 5), textcoords='offset points',
                        fontsize=7, color=DARK_RED)

    ax1.plot(x_fore, forecast_df['Predicted_Rate'], color=DARK_RED, linestyle='--',
             marker='D', label='5Q Forecast', linewidth=2, markersize=5)
    ax1.fill_between(x_fore, forecast_df['Lower_Bound_Periodic'],
                     forecast_df['Upper_Bound_Periodic'],
                     color=DARK_RED, alpha=0.08, label=f'{confidencelevel*100:.0f}% Forecast CI')
    ax1.xaxis.set_major_formatter(mticker.FuncFormatter(period_formatter))
    gs_style_ax(ax1, 'Figure 1: Periodic Recovery Rate Monitoring & Forecast')
    ax1.set_ylabel('Recovery Rate (%)', fontsize=10)
    ax1.legend(fontsize=8, loc='upper right', framealpha=0.9, facecolor='white',
              edgecolor=BORDER_GRAY)

    # ---- PANEL 2: Cumulative Recovery Rate ----
    ax2.plot(x_hist, hist_df['Cum_Recovery_Rate'], label='Actual Cumulative (%)',
             color=NAVY, marker='o', zorder=5, linewidth=2, markersize=5)
    ax2.plot(x_hist, hist_df['Cumulative_Baseline_Rate'], label='Expected Cumulative Baseline',
             color=DARK_GOLD, linestyle='--', zorder=4, linewidth=1.5)
    ax2.fill_between(x_hist, hist_df['Cumulative_Lower_Bound'],
                     hist_df['Cumulative_Upper_Bound'],
                     color=SLATE, alpha=0.12, label='Historical Range')
    ax2.plot(x_fore, forecast_df['Predicted_Cumulative_Rate'], color=DARK_RED,
             linestyle='--', marker='D', label='5Q Cumulative Forecast', linewidth=2, markersize=5)
    ax2.fill_between(x_fore, forecast_df['Lower_Bound_Cumulative'],
                     forecast_df['Upper_Bound_Cumulative'],
                     color=DARK_RED, alpha=0.08, label=f'{confidencelevel*100:.0f}% Forecast CI')

    target_rate = 85.0
    ax2.axhline(y=target_rate, color=DARK_GREEN, linestyle=':', linewidth=1.5,
                label=f'Target ({target_rate}%)')

    ax2.xaxis.set_major_formatter(mticker.FuncFormatter(period_formatter))
    gs_style_ax(ax2, 'Figure 2: Cumulative Recovery Deviation Analysis')
    ax2.set_xlabel('Recovery Period (Quarterly)', fontsize=10)
    ax2.set_ylabel('Cumulative Recovery (%)', fontsize=10)
    ax2.legend(fontsize=8, loc='lower right', framealpha=0.9, facecolor='white',
              edgecolor=BORDER_GRAY)

    # ---- PANEL 5: Recovery Efficiency & Anomaly Timeline ----
    efficiency = hist_df['Recovery_Rate'] / hist_df['Baseline_Rate'].clip(lower=0.01)
    eff_values = efficiency.values
    eff_colors = []
    for v in eff_values:
        if v >= 1.0:
            eff_colors.append(DARK_GREEN)
        elif v >= 0.9:
            eff_colors.append(DARK_GOLD)
        else:
            eff_colors.append(DARK_RED)

    ax3.bar(periods_hist, eff_values * 100, color=eff_colors, edgecolor='white',
            linewidth=0.3, zorder=3, width=0.7)
    ax3.axhline(y=100, color=DARK_TEXT, linewidth=0.8, linestyle='--', alpha=0.5)

    # Rolling average efficiency
    if len(eff_values) >= 3:
        roll_eff = pd.Series(eff_values * 100).rolling(3, min_periods=1).mean()
        ax3.plot(periods_hist, roll_eff, color=NAVY, linewidth=1.8, zorder=5,
                label='3Q Rolling Avg')

    # Mark anomalies
    anomalies = hist_df[hist_df['Anomaly']]
    if not anomalies.empty:
        ax3.scatter(anomalies['Month_Since_Start'],
                    anomalies['Recovery_Rate'] / anomalies['Baseline_Rate'].clip(lower=0.01) * 100,
                    color=DARK_RED, s=100, edgecolor='white', linewidths=1,
                    marker='X', label='Anomaly Detected', zorder=10)

    ax3.xaxis.set_major_formatter(mticker.FuncFormatter(period_formatter))
    gs_style_ax(ax3, 'Figure 5: Recovery Efficiency & Anomaly Timeline')
    ax3.set_ylabel('Recovery Efficiency (%)', fontsize=9)
    ax3.set_xlabel('Recovery Period (Quarterly)', fontsize=9)
    ax3.legend(fontsize=7, loc='lower left', framealpha=0.9, facecolor='white',
              edgecolor=BORDER_GRAY)
    # Legend for color coding
    ax3.text(0.98, 0.95, 'Green ≥100%  |  Gold 90-100%  |  Red <90%',
            transform=ax3.transAxes, fontsize=6, color=MED_GRAY,
            ha='right', va='top')

    # ---- PANEL 4: Recovery Progress Gauge ----
    ax4.axis('equal')
    current_cumulative = hist_df['Cum_Recovery_Rate'].iloc[-1]
    baseline_cumulative = hist_df['Cumulative_Baseline_Rate'].iloc[-1]
    forecast_final = forecast_df['Predicted_Cumulative_Rate'].iloc[-1]

    theta = np.linspace(0, np.pi, 100)
    r = 1.0
    x_circle = r * np.cos(theta)
    y_circle = r * np.sin(theta)
    ax4.fill_between(x_circle[:40], y_circle[:40], alpha=0.15, color=DARK_RED)
    ax4.fill_between(x_circle[40:70], y_circle[40:70], alpha=0.15, color=DARK_GOLD)
    ax4.fill_between(x_circle[70:], y_circle[70:], alpha=0.15, color=DARK_GREEN)

    angle_current = np.pi * (1 - current_cumulative / 100)
    ax4.arrow(0, 0, 0.7*np.cos(angle_current), 0.7*np.sin(angle_current),
              head_width=0.08, head_length=0.05, fc=NAVY, ec=NAVY, linewidth=2)
    angle_baseline = np.pi * (1 - baseline_cumulative / 100)
    ax4.plot([0, 0.5*np.cos(angle_baseline)], [0, 0.5*np.sin(angle_baseline)],
             color=DARK_GOLD, linestyle='--', linewidth=1.5)

    ax4.text(0, -0.3, f'{current_cumulative:.1f}%', ha='center', va='center',
            fontsize=16, fontweight='bold', color=NAVY)
    ax4.set_xlim(-1.2, 1.2)
    ax4.set_ylim(-0.5, 1.2)
    ax4.set_title('Figure 6: Recovery Progress Gauge', fontsize=13, fontweight='bold',
                  color=NAVY, loc='center')
    ax4.axis('off')

    progress_text = f'Target Achievement: {current_cumulative/85*100:.1f}%\n'
    progress_text += f'vs Baseline: {current_cumulative-baseline_cumulative:+.2f}%\n'
    progress_text += f'Forecast Final: {forecast_final:.2f}%'
    ax4.text(0.5, -0.15, progress_text, transform=ax4.transAxes, ha='center',
            fontsize=8, color=DARK_TEXT,
            bbox=dict(boxstyle='round', facecolor='#FAFAFA', edgecolor=BORDER_GRAY, alpha=0.9))

    # ---- PANEL 5: Cashflow Waterfall ----
    actual_cf = hist_df['Recovery_Rate'].values
    baseline_cf = hist_df['Baseline_Rate'].values
    deviation_cf = actual_cf - baseline_cf

    colors_cf = [DARK_GREEN if d >= 0 else DARK_RED for d in deviation_cf]
    ax5.bar(periods_hist, deviation_cf, color=colors_cf, edgecolor='white', linewidth=0.3, zorder=3)
    ax5.axhline(y=0, color=DARK_TEXT, linewidth=0.6, zorder=2)

    # Forecast deviation
    if len(forecast_df) > 0:
        fc_dev = forecast_df['Predicted_Rate'].values - forecast_df['Baseline_Rate'].values if 'Baseline_Rate' in forecast_df.columns else [0]*len(forecast_df)
        fc_x = forecast_df['Month_Since_Start'].values
        fc_colors = [STEEL_BLUE if d >= 0 else DARK_RED for d in fc_dev]
        ax5.bar(fc_x, fc_dev, color=fc_colors, edgecolor='white', linewidth=0.3, alpha=0.5, zorder=3, label='Forecast')

    ax5.xaxis.set_major_formatter(mticker.FuncFormatter(period_formatter))
    gs_style_ax(ax5, 'Figure 3: Cashflow Deviation Waterfall')
    ax5.set_ylabel('Deviation vs Baseline (%)', fontsize=9)
    ax5.legend(fontsize=7, loc='upper right', framealpha=0.9, facecolor='white',
              edgecolor=BORDER_GRAY)
    # Annotate largest deviation
    max_dev_idx = np.argmax(np.abs(deviation_cf))
    ax5.annotate(f'{deviation_cf[max_dev_idx]:+.1f}%',
                xy=(periods_hist[max_dev_idx], deviation_cf[max_dev_idx]),
                xytext=(0, 8 if deviation_cf[max_dev_idx] >= 0 else -12), textcoords='offset points',
                fontsize=7, color=DARK_RED if deviation_cf[max_dev_idx] < 0 else DARK_GREEN,
                ha='center', fontweight='bold')

    # ---- PANEL 6: Deviation Trend & Volatility ----
    cum_dev = (hist_df['Cum_Recovery_Rate'] - hist_df['Cumulative_Baseline_Rate']).values
    ax6.plot(periods_hist, cum_dev, color=NAVY, linewidth=2, marker='o', markersize=5, label='Cumulative Deviation', zorder=5)
    ax6.fill_between(periods_hist, 0, cum_dev, color=NAVY, alpha=0.08)

    # Rolling volatility of deviation
    if len(periods_hist) >= 4:
        rolling_std = pd.Series(deviation_cf).rolling(4, min_periods=2).std().values
        ax6_twin = ax6.twinx()
        ax6_twin.plot(periods_hist, rolling_std, color=DARK_GOLD, linewidth=1.2,
                     linestyle='--', label='Deviation Volatility (4Q rolling)')
        ax6_twin.set_ylabel('Volatility (σ)', fontsize=8, color=DARK_GOLD)
        ax6_twin.tick_params(axis='y', colors=DARK_GOLD, labelsize=7)
        ax6_twin.legend(fontsize=7, loc='upper left', framealpha=0.9, facecolor='white',
                       edgecolor=BORDER_GRAY)
        ax6_twin.spines['right'].set_color(DARK_GOLD)
        ax6_twin.spines['top'].set_visible(False)

    ax6.axhline(y=0, color=MED_GRAY, linewidth=0.5, linestyle=':')
    ax6.xaxis.set_major_formatter(mticker.FuncFormatter(period_formatter))
    gs_style_ax(ax6, 'Figure 4: Cumulative Deviation Trend & Risk Drift')
    ax6.set_ylabel('Cumulative Deviation (%)', fontsize=9)
    ax6.set_xlabel('Recovery Period (Quarterly)', fontsize=9)
    ax6.legend(fontsize=7, loc='lower left', framealpha=0.9, facecolor='white',
              edgecolor=BORDER_GRAY)

    # Mark deviation trajectory
    if len(periods_hist) >= 2:
        recent_trend = 'improving' if cum_dev[-1] > cum_dev[-2] else 'deteriorating'
        trend_color = DARK_GREEN if recent_trend == 'improving' else DARK_RED
        ax6.annotate(f'Recent: {recent_trend}',
                    xy=(periods_hist[-1], cum_dev[-1]),
                    xytext=(10, 10), textcoords='offset points',
                    fontsize=7, color=trend_color, fontweight='bold',
                    arrowprops=dict(arrowstyle='->', color=trend_color, alpha=0.6))

    current_period = hist_df['Month_Since_Start'].iloc[-1]
    fig.suptitle(f'[{project_name}]  Post-Investment Risk Assessment  |  Period {current_period}',
                 fontsize=18, fontweight='bold', color=NAVY, y=0.98)

    # Footer
    fig.text(0.5, 0.01, 'Source: NPL ABS Quant Risk Desk | CONFIDENTIAL',
             fontsize=8, color=MED_GRAY, ha='center')

    fig_path = f"npl_abs_gs_analysis_chart.png"
    plt.savefig(fig_path, dpi=150, facecolor='white', edgecolor='none')
    print(f"[GS Chart] Saved: {fig_path}")
    plt.close(fig)
    return fig_path

def build_report_prompt(hist_df, forecast_df, project_name, bank, risk_rating, stress_results):
    anomalies = hist_df[hist_df['Anomaly']]
    last_hist_period = hist_df.iloc[-1]
    last_fore_period = forecast_df.iloc[-1]
    rating_map = {"正常类": 0, "关注类": 1, "次级类": 2, "可疑类": 3}
    report_date = datetime.now().strftime('%Y-%m-%d')

    # 判断项目所处阶段
    current_period = len(hist_df)
    if current_period <= 4:
        project_stage = "早期"
        stage_focus = "数据质量监控、趋势初判、基准校准"
    elif current_period <= 12:
        project_stage = "中期"
        stage_focus = "司法处置效率、异常点根因、回收节奏"
    else:
        project_stage = "晚期"
        stage_focus = "最终回收达成率、尾部风险、核销准备"

    # 计算关键对比指标
    same_bank_benchmark = last_hist_period.get('Same_Bank_Cumulative_Peer_Rate', None)
    current_cumulative = last_hist_period['Cum_Recovery_Rate']

    # 同银行排名估算
    if same_bank_benchmark:
        vs_same_bank = current_cumulative - same_bank_benchmark
        if vs_same_bank > 2:
            peer_ranking = "优于同银行平均水平"
        elif vs_same_bank < -2:
            peer_ranking = "落后同银行平均水平"
        else:
            peer_ranking = "与同银行平均水平持平"
        same_bank_comparison = f"- 当前累计回收率: {current_cumulative:.2f}%\n- 同银行平均水平: {same_bank_benchmark:.2f}%\n- 偏离度: {vs_same_bank:+.2f}%"
    else:
        peer_ranking = "同银行数据不足"
        same_bank_comparison = ""

    # 构建异常点详情
    anomaly_details = ""
    if not anomalies.empty:
        for _, row in anomalies.iterrows():
            period = int(row['Month_Since_Start']) + 1
            actual = row['Recovery_Rate']
            baseline = row['Baseline_Rate']
            direction = "高于" if actual > baseline else "低于"
            anomaly_details += f"回收期 {period}: 实际回收率 {actual:.2f}%，{direction}预期基准({baseline:.2f}%)"

    # 计算偏离度
    deviation_from_lower = last_hist_period['Cumulative_Lower_Bound'] - last_hist_period['Cum_Recovery_Rate']
    deviation_direction = "高" if deviation_from_lower < 0 else "低"
    deviation_abs = abs(deviation_from_lower)

    # 构建未来5期预测表格
    forecast_table = ""
    for _, row in forecast_df.iterrows():
        period_num = int(row['Month_Since_Start']) + 1
        date_str = row['Date'].strftime('%Y-%m-%d') if pd.notna(row['Date']) else 'N/A'
        periodic_rate = row['Predicted_Rate']
        cumulative_rate = row['Predicted_Cumulative_Rate']
        cum_lower = row['Lower_Bound_Cumulative']
        cum_upper = row['Upper_Bound_Cumulative']
        forecast_table += f"第{period_num}期({date_str}): 当期{periodic_rate:.2f}%, 累计{cumulative_rate:.2f}% (区间:{cum_lower:.2f}%-{cum_upper:.2f}%)\n"

    # 构建压力测试结果
    if not stress_results:
        stress_test_results = "结果: 所有压力情景下评级保持稳定。"
    else:
        stress_test_results = "评级下调临界点:"
        for rating, factor in sorted(stress_results.items(), key=lambda item: rating_map[item[0]]):
            stress_test_results += f"综合压力因子降至 {factor:.2f} 时，评级下调至 {rating}"

    # 读取模板并填充数据
    template_path = os.path.join(os.path.dirname(__file__), 'templates', 'prompt.txt')
    try:
        with open(template_path, 'r', encoding='utf-8') as f:
            template = f.read()
    except FileNotFoundError:
        raise FileNotFoundError(f"模板文件未找到: {template_path}")

    briefing = template.format(
        project_name=project_name,
        bank=bank,
        report_date=report_date,
        risk_rating=risk_rating,
        project_stage=project_stage,
        current_period=current_period,
        stage_focus=stage_focus,
        peer_ranking=peer_ranking,
        same_bank_comparison=same_bank_comparison,
        anomaly_count=len(anomalies),
        anomaly_details=anomaly_details,
        current_cumulative=current_cumulative,
        baseline_rate=last_hist_period['Cumulative_Baseline_Rate'],
        lower_bound=last_hist_period['Cumulative_Lower_Bound'],
        deviation_direction=deviation_direction,
        deviation_abs=deviation_abs,
        forecast_table=forecast_table,
        final_cumulative_rate=last_fore_period['Predicted_Cumulative_Rate'],
        final_lower_bound=last_fore_period['Lower_Bound_Cumulative'],
        final_upper_bound=last_fore_period['Upper_Bound_Cumulative'],
        confidence_level=CONFIDENCE_LEVEL * 100,
        stress_test_results=stress_test_results
    )

    return briefing

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
def generate_ai_report(api_key, briefing_text):
    print("正在调用AI大模型生成报告")
    url = "https://api.deepseek.com/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}

    payload = {
        "model": "deepseek-v4-flash",
        "messages": [{"role": "user", "content": briefing_text}],
        "response_format": {"type": "json_object"},
        "stream": False
    }

    response = requests.post(
        url,
        headers=headers,
        json=payload,
        timeout=90,
        proxies={"http": None, "https": None}
    )
    response.raise_for_status()

    result = response.json()
    report_text = result['choices'][0]['message']['content'].strip()

    if report_text.startswith("```json"):
        report_text = report_text[7:]
    if report_text.endswith("```"):
        report_text = report_text[:-3]
    report_text = report_text.strip()

    try:
        report_dict = json.loads(report_text)
        print("JSON解析成功")
        return report_dict
    except json.JSONDecodeError as e:
        print(f"返回内容非合法JSON，准备触发自动重试。{report_text[:100]}")
        raise Exception(f"JSON解析失败: {e}")

def export_raw_json(report_dict, project_name):
    try:
        raw_path = f"{project_name}_RAW_AI_Output.json"
        with open(raw_path, 'w', encoding='utf-8') as f:
            json.dump(report_dict, f, ensure_ascii=False, indent=4)
        print(f"原始JSON数据已保存至:{raw_path}")
    except Exception as e:
        print(f"原始JSON保存失败:{e}")

def export_to_word_from_json(report_dict, project_name, fig_path, risk_metrics=None):
    """
    Goldman Sachs Research Report Style — Word Export
    封面页 | KPI仪表板 | 正文 | 图表 | 页眉页脚
    Color: Navy #003A70 / Dark Gold #C4A43E / White
    """
    from datetime import datetime
    try:
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Arial'
        style.font.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.size = Pt(10)

        # === Goldman Sachs Color Palette ===
        NAVY = RGBColor(0x00, 0x3A, 0x70)
        DARK_GOLD = RGBColor(0xC4, 0xA4, 0x3E)
        DARK_TEXT = RGBColor(0x33, 0x33, 0x33)
        MED_GRAY = RGBColor(0x99, 0x99, 0x99)
        LIGHT_GRAY_BG = RGBColor(0xF5, 0xF5, 0xF5)
        DARK_GREEN = RGBColor(0x1E, 0x84, 0x49)
        DARK_RED = RGBColor(0xC0, 0x39, 0x2B)
        DARK_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
        rating_colors = {
            "正常类": RGBColor(0x1E, 0x84, 0x49),
            "关注类": RGBColor(0xC4, 0xA4, 0x3E),
            "次级类": RGBColor(0xE6, 0x7E, 0x22),
            "可疑类": RGBColor(0xC0, 0x39, 0x2B)
        }

        def set_run(run, size=10, color=DARK_TEXT, bold=False):
            """Helper: set CJK-compatible font on a run."""
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold

        def add_hr(doc, color=DARK_GOLD, width_pt=1.5, space_after=8):
            """Thin horizontal rule via paragraph bottom border."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(space_after)
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): str(int(width_pt * 8)),
                qn('w:space'): '1',
                qn('w:color'): str(color).replace('#', ''),
            })
            pBdr.append(bottom)
            pPr.append(pBdr)

        # ============================================================
        # COVER PAGE
        # ============================================================
        for _ in range(6):
            document.add_paragraph()

        # Title
        p_title = document.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_title.add_run('NPL ABS 投后风险评估报告')
        set_run(r, size=26, color=NAVY, bold=True)
        p_title.paragraph_format.space_after = Pt(8)

        # English subtitle
        p_sub = document.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_sub.add_run('Post-Investment Risk Assessment Report')
        set_run(r, size=14, color=DARK_GOLD)
        p_sub.paragraph_format.space_after = Pt(24)

        # Horizontal rule
        add_hr(document, DARK_GOLD, width_pt=1.5, space_after=20)

        # Project name
        p_name = document.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_name.add_run(project_name)
        set_run(r, size=18, color=DARK_TEXT, bold=True)
        p_name.paragraph_format.space_after = Pt(8)

        # Date
        p_date = document.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_date.add_run(f"报告日期: {datetime.now().strftime('%Y年%m月%d日')}")
        set_run(r, size=11, color=MED_GRAY)
        p_date.paragraph_format.space_after = Pt(20)

        # Risk rating badge
        rating_text = report_dict.get('risk_rating', '')
        if rating_text:
            rating_value = rating_text.split('：')[-1] if '：' in rating_text else rating_text
            p_rating = document.add_paragraph()
            p_rating.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p_rating.add_run('综合风险评级  ')
            set_run(r, size=12, color=MED_GRAY)
            rating_color = DARK_TEXT
            for key, color in rating_colors.items():
                if key in rating_text:
                    rating_color = color
                    break
            r = p_rating.add_run(rating_value)
            set_run(r, size=16, color=rating_color, bold=True)

        # Bottom rule
        add_hr(document, NAVY, width_pt=0.5, space_after=12)

        # Confidentiality notice
        p_disc = document.add_paragraph()
        p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_disc.add_run('CONFIDENTIAL — 本报告为量化模型自动生成，仅供内部参考')
        set_run(r, size=7, color=MED_GRAY)

        # ============================================================
        # PAGE 2: KPI DASHBOARD
        # ============================================================
        document.add_page_break()
        h_kpi = document.add_heading('关键指标仪表板', level=1)
        for r in h_kpi.runs:
            set_run(r, size=16, color=NAVY, bold=True)

        rating_text = report_dict.get('risk_rating', 'N/A')
        risklevel = rating_text.split('：')[-1] if '：' in rating_text else rating_text
        risklevel = risklevel.split('（')[0].split('(')[0].strip()  # drop parenthetical
        trend_text = str(risk_metrics.trend.value) if risk_metrics and hasattr(risk_metrics, 'trend') else 'N/A'
        confidence_text = f"{risk_metrics.rating_confidence:.1%}" if risk_metrics else 'N/A'
        max_dd_text = f"{risk_metrics.max_drawdown:.2f}%" if risk_metrics else 'N/A'

        # 2x2 KPI table
        table = document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        kpi_data = [
            [
                ('风险评级', risklevel, rating_colors.get(risklevel, DARK_TEXT)),
                ('评级置信度', confidence_text, DARK_TEXT),
            ],
            [
                ('趋势判断', trend_text, DARK_TEXT),
                ('最大回撤', max_dd_text, DARK_TEXT),
            ]
        ]
        for i, row_data in enumerate(kpi_data):
            row = table.rows[i]
            for j, (label, value, val_color) in enumerate(row_data):
                cell = row.cells[j]
                cell.text = ''
                p1 = cell.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p1.add_run(label)
                set_run(r, size=9, color=MED_GRAY)
                p2 = cell.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p2.add_run(value)
                set_run(r, size=20, color=val_color, bold=True)

        document.add_paragraph()

        # ---- Key Observations (auto-generated) ----
        h_obs = document.add_heading('关键指标解读', level=2)
        for r in h_obs.runs:
            set_run(r, size=13, color=DARK_GOLD, bold=True)

        risk_rating_text = report_dict.get('risk_rating', '')
        observations = []
        trend_val = str(risk_metrics.trend.value) if risk_metrics and hasattr(risk_metrics, 'trend') else ''
        trend_map = {
            'IMPROVING': ('改善中', DARK_GREEN, '回收表现持续向好，当前策略有效'),
            'STABLE': ('稳定', DARK_TEXT, '回收进度符合预期，无显著偏离'),
            'DETERIORATING': ('恶化中', DARK_RED, '需重点关注，建议排查底层资产质量变化'),
            'VOLATILE': ('波动较大', DARK_ORANGE, '回收率离散度偏高，可能存在季节性/结构性因素'),
        }
        if trend_val in trend_map:
            label, color, advice = trend_map[trend_val]
            observations.append((f'趋势判断：{label}', color, advice))
        if risk_metrics:
            dd = risk_metrics.max_drawdown
            if dd < 2:
                observations.append((f'最大回撤：低风险 ({dd:.1f}%)', DARK_GREEN, '处于安全区间'))
            elif dd < 5:
                observations.append((f'最大回撤：中等风险 ({dd:.1f}%)', DARK_GOLD, '建议持续监控'))
            else:
                observations.append((f'最大回撤：高风险 ({dd:.1f}%)', DARK_RED, '接近预警阈值'))
            conf = risk_metrics.rating_confidence
            if conf > 0.8:
                observations.append((f'评级置信度：高 ({conf:.0%})', DARK_GREEN, '评级结果可靠'))
            elif conf > 0.6:
                observations.append((f'评级置信度：中 ({conf:.0%})', DARK_GOLD, '建议结合人工判断'))
            else:
                observations.append((f'评级置信度：低 ({conf:.0%})', DARK_RED, '数据质量或样本量可能不足'))
        if '正常' in risk_rating_text:
            observations.append(('综合评估：正常类', DARK_GREEN, '当前风险可控，按常规频率跟踪即可'))
        elif '关注' in risk_rating_text:
            observations.append(('综合评估：关注类', DARK_GOLD, '存在若干待关注指标，建议提高监测频率'))
        elif '次级' in risk_rating_text:
            observations.append(('综合评估：次级类', DARK_ORANGE, '风险显著上升，需制定应对预案'))
        elif '可疑' in risk_rating_text:
            observations.append(('综合评估：可疑类', DARK_RED, '风险处于高位，建议立即启动应急处置'))

        for title, color, desc in observations:
            p = document.add_paragraph()
            r1 = p.add_run(f'■ {title}')
            set_run(r1, size=9, color=color, bold=True)
            r2 = p.add_run(f' — {desc}')
            set_run(r2, size=9, color=DARK_TEXT)

        document.add_paragraph()

        # Executive summary
        h_summary = document.add_heading('核心结论', level=1)
        for r in h_summary.runs:
            set_run(r, size=16, color=NAVY, bold=True)

        for summary_point in report_dict.get('executive_summary', []):
            p = document.add_paragraph(summary_point, style='List Bullet')
            p.paragraph_format.left_indent = Pt(20)
            for r in p.runs:
                set_run(r, size=10, color=DARK_TEXT)

        # ============================================================
        # PAGE 3+: CHARTS & SECTIONS
        # ============================================================
        document.add_page_break()
        if fig_path and os.path.exists(fig_path):
            h_chart = document.add_heading('图表分析', level=1)
            for r in h_chart.runs:
                set_run(r, size=16, color=NAVY, bold=True)

            document.add_picture(fig_path, width=Pt(480))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = caption.add_run(f'图1: {project_name} 投后风险分析图表')
            set_run(r, size=9, color=MED_GRAY)
            document.add_page_break()

        for section in report_dict.get('sections', []):
            h_sec = document.add_heading(section.get('section_title', ''), level=1)
            for r in h_sec.runs:
                set_run(r, size=16, color=NAVY, bold=True)

            for sub in section.get('subsections', []):
                h_sub = document.add_heading(sub.get('sub_title', ''), level=2)
                for r in h_sub.runs:
                    set_run(r, size=13, color=DARK_GOLD, bold=True)

                for paragraph_text in sub.get('content', []):
                    p = document.add_paragraph(paragraph_text)
                    p.paragraph_format.first_line_indent = Pt(24)
                    p.paragraph_format.line_spacing = 1.5
                    for r in p.runs:
                        set_run(r, size=10, color=DARK_TEXT)

        # ============================================================
        # HEADER & FOOTER
        # ============================================================
        section = document.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run(f"NPL ABS 投后风险评估  |  {project_name}")
        set_run(r, size=8, color=MED_GRAY)
        # Navy accent line under header
        hpPr = hp._p.get_or_add_pPr()
        pBdr = hpPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '4',
            qn('w:color'): '003A70',
        })
        pBdr.append(bottom)
        hpPr.append(pBdr)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("— CONFIDENTIAL —")
        set_run(r, size=7, color=MED_GRAY)

        # ============================================================
        # SAVE
        # ============================================================
        doc_path = f"{project_name}_V16_GS_Report.docx"
        document.save(doc_path)
        print(f"报告已保存至: {doc_path}")

        try:
            if os.name == 'nt':
                os.startfile(doc_path)
            elif os.name == 'posix':
                os.system(f'open "{doc_path}"')
        except:
            pass

    except Exception as e:
        import traceback
        print(f"文档导出失败:{e}")
        traceback.print_exc()


def export_to_pdf_from_json(report_dict, project_name, fig_path, risk_metrics=None):
    """
    Goldman Sachs Research Style — PDF Export (ReportLab)
    Simple layout: Cover | KPI | Chart | Sections
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image, Flowable,
    )
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from PIL import Image as PILImage
    import os as _os

    # === CJK Font Registration ===
    _font_dir = _os.environ.get('WINDIR', '') + r'\Fonts'
    for _fn, _fname, _idx in [('YaHei', 'msyh.ttc', 0), ('YaHeiBold', 'msyhbd.ttc', 0)]:
        _fp = _os.path.join(_font_dir, _fname)
        if _os.path.isfile(_fp):
            try:
                pdfmetrics.registerFont(TTFont(_fn, _fp, subfontIndex=_idx))
            except:
                pass

    # === Colors ===
    NAVY = HexColor('#003A70')
    DARK_GOLD = HexColor('#C4A43E')
    DARK_TEXT = HexColor('#333333')
    MED_GRAY = HexColor('#999999')
    LIGHT_GRAY = HexColor('#E0E0E0')
    WHITE = HexColor('#FFFFFF')
    OFF_WHITE = HexColor('#FAFAFA')
    DARK_RED = HexColor('#C0392B')
    DARK_GREEN = HexColor('#1E8449')

    PAGE_W, PAGE_H = A4
    page_width = PAGE_W - 70

    # === Styles ===
    styles = getSampleStyleSheet()

    body_style = ParagraphStyle('BodyGS', parent=styles['Normal'],
        fontName='YaHei', fontSize=9, leading=16, textColor=DARK_TEXT,
        spaceAfter=9, alignment=TA_JUSTIFY)

    h1_style = ParagraphStyle('H1GS', parent=styles['Heading1'],
        fontName='YaHeiBold', fontSize=16, textColor=NAVY,
        spaceAfter=10, spaceBefore=18)

    h2_style = ParagraphStyle('H2GS', parent=styles['Heading2'],
        fontName='YaHeiBold', fontSize=13, textColor=DARK_GOLD,
        spaceAfter=8, spaceBefore=14)

    caption_style = ParagraphStyle('CaptionGS',
        fontName='YaHei', fontSize=7, textColor=MED_GRAY,
        alignment=TA_CENTER, spaceAfter=6)

    # === Helpers ===
    class HRule(Flowable):
        def __init__(self, width, color=DARK_GOLD, thickness=1.5):
            Flowable.__init__(self)
            self.width = width
            self.height = thickness + 8
            self._color = color
            self._thickness = thickness
        def draw(self):
            self.canv.setStrokeColor(self._color)
            self.canv.setLineWidth(self._thickness)
            self.canv.line(0, 4, self.width, 4)

    def S(h=8):
        return Spacer(1, h)
    def hr(color=DARK_GOLD, t=1.5):
        return HRule(page_width, color, t)
    def H1(text):
        return Paragraph(text, h1_style)
    def H2(text):
        return Paragraph(text, h2_style)
    def B(text):
        return Paragraph(text, body_style)

    # === Page Template ===
    class GSCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            canvas.Canvas.__init__(self, *args, **kwargs)
            self._saved = []
        def showPage(self):
            self._saved.append(dict(self.__dict__))
            self._startPage()
        def save(self):
            for s in self._saved:
                self.__dict__.update(s)
                n = self._pageNumber
                # Top accent line
                self.setStrokeColor(DARK_GOLD)
                self.setLineWidth(1.5)
                self.line(30, PAGE_H - 35, PAGE_W - 30, PAGE_H - 35)
                # Footer
                self.setFont('YaHei', 7)
                self.setFillColor(MED_GRAY)
                self.drawRightString(PAGE_W - 30, 25, f"Page {n} / {len(self._saved)}")
                self.drawString(30, 25, "CONFIDENTIAL — NPL ABS Quant Risk Desk")
                # Header logo
                self.setFillColor(NAVY)
                self.setFont('YaHeiBold', 8)
                self.drawString(30, PAGE_H - 28, f"NPL ABS 投后风险评估 | {project_name}")
                canvas.Canvas.showPage(self)
            canvas.Canvas.save(self)

    # === Build ===
    doc = SimpleDocTemplate(
        f"{project_name}_V16_GS_Report.pdf",
        pagesize=A4, leftMargin=35, rightMargin=35,
        topMargin=40, bottomMargin=35,
        title=f'NPL ABS Report — {project_name}',
        author='NPL ABS Quant Risk Desk',
    )
    story = []

    # ---- COVER PAGE ----
    story.append(S(110))
    story.append(Paragraph('NPL ABS 投后风险评估报告', ParagraphStyle('T',
        fontName='YaHeiBold', fontSize=26, textColor=NAVY,
        alignment=TA_CENTER, spaceAfter=20, leading=34)))
    story.append(Paragraph('Post-Investment Risk Assessment Report', ParagraphStyle('ST',
        fontName='YaHei', fontSize=14, textColor=DARK_GOLD,
        alignment=TA_CENTER, spaceAfter=28, leading=18)))
    story.append(hr(DARK_GOLD, 1.5))
    story.append(S(30))
    story.append(Paragraph(project_name, ParagraphStyle('PN',
        fontName='YaHeiBold', fontSize=18, textColor=DARK_TEXT,
        alignment=TA_CENTER, spaceAfter=12, leading=24)))
    story.append(Paragraph(f'报告日期: {datetime.now().strftime("%Y年%m月%d日")}',
        ParagraphStyle('D', fontName='YaHei', fontSize=11, textColor=MED_GRAY,
        alignment=TA_CENTER, spaceAfter=24, leading=14)))

    # Rating
    rating_text = report_dict.get('risk_rating', '')
    if rating_text:
        rv = rating_text.split('：')[-1] if '：' in rating_text else rating_text
        rv = rv.split('（')[0].split('(')[0].strip()  # drop parenthetical
        rc = DARK_TEXT
        for k, c in [("正常类", DARK_GREEN), ("关注类", DARK_GOLD),
                      ("次级类", HexColor('#E67E22')), ("可疑类", DARK_RED)]:
            if k in rating_text:
                rc = c
                break
        story.append(Paragraph(f'综合风险评级: <font color="{rc}"><b>{rv}</b></font>',
            ParagraphStyle('R', fontName='YaHei', fontSize=12, textColor=MED_GRAY,
            alignment=TA_CENTER, spaceAfter=16)))
    story.append(hr(NAVY, 0.5))
    story.append(S(20))
    story.append(Paragraph('CONFIDENTIAL — 本报告为量化模型自动生成，仅供内部参考',
        ParagraphStyle('DS', fontName='YaHei', fontSize=7, textColor=MED_GRAY,
        alignment=TA_CENTER)))

    # ---- PAGE 2: KPI ----
    story.append(PageBreak())
    story.append(H1('关键指标仪表板'))
    story.append(hr(DARK_GOLD, 0.5))
    story.append(S(12))

    rt = report_dict.get('risk_rating', 'N/A')
    rl = rt.split("：")[-1] if "：" in rt else rt
    rl = rl.split("（")[0].split("(")[0].strip()  # short label only
    tr = str(risk_metrics.trend.value) if risk_metrics and hasattr(risk_metrics, 'trend') else 'N/A'
    cf = f"{risk_metrics.rating_confidence:.1%}" if risk_metrics else 'N/A'
    md = f"{risk_metrics.max_drawdown:.2f}%" if risk_metrics else 'N/A'

    kt = Table([
        ['风险评级', rl, '评级置信度', cf],
        ['趋势判断', tr, '最大回撤', md],
    ], colWidths=[page_width*0.2, page_width*0.3, page_width*0.2, page_width*0.3])
    kt.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'YaHei'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('FONTNAME', (1, 0), (1, -1), 'YaHeiBold'),
        ('FONTNAME', (3, 0), (3, -1), 'YaHeiBold'),
        ('FONTSIZE', (1, 0), (1, -1), 14),
        ('FONTSIZE', (3, 0), (3, -1), 14),
        ('TEXTCOLOR', (0, 0), (-1, -1), DARK_TEXT),
        ('TEXTCOLOR', (0, 0), (0, -1), MED_GRAY),
        ('TEXTCOLOR', (2, 0), (2, -1), MED_GRAY),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.4, LIGHT_GRAY),
        ('ROWBACKGROUNDS', (0, 0), (-1, -1), [OFF_WHITE, WHITE]),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(kt)
    story.append(S(16))

    # ---- Auto-generated Key Observations ----
    story.append(H2('关键指标解读'))
    observations = []
    # Trend commentary
    trend_val = str(risk_metrics.trend.value) if risk_metrics and hasattr(risk_metrics, 'trend') else ''
    trend_map = {
        'IMPROVING': ('改善中', DARK_GREEN, '回收表现持续向好，当前策略有效'),
        'STABLE': ('稳定', DARK_TEXT, '回收进度符合预期，无显著偏离'),
        'DETERIORATING': ('恶化中', DARK_RED, '需重点关注，建议排查底层资产质量变化'),
        'VOLATILE': ('波动较大', HexColor('#E67E22'), '回收率离散度偏高，可能存在季节性/结构性因素'),
    }
    if trend_val in trend_map:
        label, color, advice = trend_map[trend_val]
        observations.append((f'趋势判断：{label}', color, advice))

    # Max drawdown commentary
    if risk_metrics:
        dd = risk_metrics.max_drawdown
        if dd < 2:
            observations.append(('最大回撤：低风险', DARK_GREEN, f'当前最大回撤 {dd:.1f}%，处于安全区间'))
        elif dd < 5:
            observations.append(('最大回撤：中等风险', DARK_GOLD, f'当前最大回撤 {dd:.1f}%，建议持续监控'))
        else:
            observations.append(('最大回撤：高风险', DARK_RED, f'当前最大回撤 {dd:.1f}%，接近预警阈值'))

    # Rating confidence
    if risk_metrics:
        conf = risk_metrics.rating_confidence
        if conf > 0.8:
            observations.append(('评级置信度：高', DARK_GREEN, f'模型置信度 {conf:.0%}，评级结果可靠'))
        elif conf > 0.6:
            observations.append(('评级置信度：中', DARK_GOLD, f'模型置信度 {conf:.0%}，建议结合人工判断'))
        else:
            observations.append(('评级置信度：低', DARK_RED, f'模型置信度 {conf:.0%}，数据质量或样本量可能不足'))

    # Risk rating summary
    rt = report_dict.get('risk_rating', '')
    if '正常' in rt:
        observations.append(('综合评估：正常类', DARK_GREEN, '当前风险可控，按常规频率跟踪即可'))
    elif '关注' in rt:
        observations.append(('综合评估：关注类', DARK_GOLD, '存在若干待关注指标，建议提高监测频率'))
    elif '次级' in rt:
        observations.append(('综合评估：次级类', HexColor('#E67E22'), '风险显著上升，需制定应对预案'))
    elif '可疑' in rt:
        observations.append(('综合评估：可疑类', DARK_RED, '风险处于高位，建议立即启动应急处置'))

    for title, color, desc in observations:
        story.append(Paragraph(
            f'<font color="{color}"><b>■ {title}</b></font> — {desc}',
            ParagraphStyle('Obs', fontName='YaHei', fontSize=9, textColor=DARK_TEXT,
                          leftIndent=10, spaceAfter=6, leading=14)))

    story.append(S(10))
    story.append(H2('核心结论'))
    for sp in report_dict.get('executive_summary', []):
        story.append(Paragraph(f'• {sp}', ParagraphStyle('BL',
            fontName='YaHei', fontSize=9, textColor=DARK_TEXT,
            leftIndent=20, spaceAfter=6)))

    # ---- PAGE 3: CHART ----
    story.append(PageBreak())
    if fig_path and _os.path.exists(fig_path):
        story.append(H1('图表分析'))
        story.append(hr(DARK_GOLD, 0.5))
        story.append(S(10))
        # Read actual image dimensions with PIL, calculate correct height
        try:
            pil_img = PILImage.open(fig_path)
            pw, ph = pil_img.size
            img_w = page_width * 0.85
            img_h = img_w * ph / pw
            pil_img.close()
        except:
            img_w = page_width * 0.85
            img_h = img_w * 0.8  # fallback 5:4
        story.append(Image(_os.path.abspath(fig_path), width=img_w, height=img_h))
        story.append(Paragraph(f'图1: {project_name} 投后风险分析图表', caption_style))
    story.append(PageBreak())

    # ---- SECTIONS ----
    for section in report_dict.get('sections', []):
        story.append(H1(section.get('section_title', '')))
        story.append(hr(DARK_GOLD, 0.5))
        story.append(S(6))
        for sub in section.get('subsections', []):
            story.append(H2(sub.get('sub_title', '')))
            for pt in sub.get('content', []):
                story.append(B(pt))
                story.append(S(4))

    # ---- BUILD ----
    doc.build(story, canvasmaker=GSCanvas)
    pdf_path = f"{project_name}_V16_GS_Report.pdf"
    print(f"[GS PDF] Report saved: {pdf_path}")
    try:
        if _os.name == 'nt':
            _os.startfile(pdf_path)
    except:
        pass


# ProjectAnalyzer类
class RiskTrend(Enum):
    IMPROVING = "改善"
    STABLE = "稳定"
    DETERIORATING = "恶化"
    VOLATILE = "波动"

@dataclass
class RiskMetrics:
    periodic_volatility: float
    cumulative_volatility: float
    tail_risk_95: float
    max_drawdown: float
    trend: 'RiskTrend'
    trend_strength: float
    rating_confidence: float

class VolatilityEstimator:
    """鲁棒的波动率估计器"""

    @staticmethod
    def mad(series: pd.Series) -> float:
        """中位数绝对偏差 - 对异常值鲁棒"""
        median = series.median()
        mad_val = (series - median).abs().median()
        return mad_val / 0.6745 if mad_val > 0 else 0.01

    @staticmethod
    def ewma_volatility(series: pd.Series, lambda_factor: float = 0.94) -> pd.Series:
        """EWMA波动率 - 给近期数据更高权重"""
        squared_returns = series ** 2
        ewma_var = squared_returns.ewm(alpha=1-lambda_factor, adjust=False).mean()
        return np.sqrt(ewma_var)

    @staticmethod
    def rolling_realized_vol(series: pd.Series, window: int = 4) -> pd.Series:
        """滚动实现波动率"""
        return series.rolling(window=window, min_periods=2).std()

    @classmethod
    def estimate(cls, deviation_series: pd.Series, method: str = 'adaptive') -> pd.Series:
        """自适应波动率估计"""
        n = len(deviation_series.dropna())

        if method == 'mad' or n < 8:
            vol = pd.Series(cls.mad(deviation_series), index=deviation_series.index)
        elif method == 'ewma' or n < 16:
            vol = cls.ewma_volatility(deviation_series)
        else:
            ewma_vol = cls.ewma_volatility(deviation_series)
            rolling_vol = cls.rolling_realized_vol(deviation_series)
            vol = 0.7 * ewma_vol + 0.3 * rolling_vol

        return vol.fillna(vol.dropna().iloc[0] if len(vol.dropna()) > 0 else 0.01)

class ConfidenceIntervalCalculator:
    """非参数置信区间计算"""

    def __init__(self, confidencelevel: float = 0.99):
        self.confidencelevel = confidencelevel

    def calculate_bootstrap_ci(self, baseline: pd.Series, historical_deviations: pd.Series,
                               n_bootstrap: int = 1000) -> Tuple[pd.Series, pd.Series]:
        """Bootstrap置信区间"""
        np.random.seed(42)
        deviations_clean = historical_deviations.dropna()

        if len(deviations_clean) < 5:
            std = deviations_clean.std() if len(deviations_clean) > 0 else 0.02
            z = stats.norm.ppf(1 - (1 - self.confidencelevel) / 2)
            return baseline - z * std, baseline + z * std

        bootstrap_samples = np.random.choice(deviations_clean, size=(n_bootstrap, len(baseline)), replace=True)
        alpha = 1 - self.confidencelevel
        lower_percentile = alpha / 2 * 100
        upper_percentile = (1 - alpha / 2) * 100
        lower_bounds = baseline + np.percentile(bootstrap_samples, lower_percentile, axis=0)
        upper_bounds = baseline + np.percentile(bootstrap_samples, upper_percentile, axis=0)
        return pd.Series(lower_bounds, index=baseline.index), pd.Series(upper_bounds, index=baseline.index)

@dataclass
class DynamicThresholds:
    """动态评级阈值"""
    normal_to_watch: float
    watch_to_subprime: float
    subprime_to_doubtful: float

    @classmethod
    def calculate(cls, df: pd.DataFrame, data_quality_score: float):
        """基于项目特征动态计算阈值"""
        n_periods = len(df)
        volatility = df['Deviation_Rate'].std() if 'Deviation_Rate' in df.columns else 0.05

        base_normal = 1.0
        base_subprime = 5.0
        base_doubtful = 10.0

        quality_factor = 0.8 + 0.4 * (data_quality_score / 100)
        progress = df['Cum_Recovery_Rate'].iloc[-1] / 100 if 'Cum_Recovery_Rate' in df.columns else 0.5
        if progress < 0.3:
            stage_factor = 1.3
        elif progress < 0.7:
            stage_factor = 1.0
        else:
            stage_factor = 0.8

        vol_factor = 1 + (volatility - 0.02) * 10
        vol_factor = np.clip(vol_factor, 0.8, 1.5)

        return cls(
            normal_to_watch=base_normal * quality_factor * stage_factor * vol_factor,
            watch_to_subprime=base_subprime * quality_factor * stage_factor * vol_factor,
            subprime_to_doubtful=base_doubtful * quality_factor * stage_factor * vol_factor
        )

class TrendAnalyzer:
    """趋势分析器"""

    @staticmethod
    def analyze_deviationtrend(df: pd.DataFrame, window: int = 4) -> Tuple['RiskTrend', float]:
        """分析偏离度趋势"""
        if len(df) < window + 1:
            return RiskTrend.STABLE, 0.5

        recent = df.tail(window)
        deviations = recent['Cum_Deviation_Rate'] if 'Cum_Deviation_Rate' in recent.columns else recent['Deviation_Rate']
        x = np.arange(len(deviations))
        coeffs = np.polyfit(x, deviations, 1)
        slope = coeffs[0]
        predicted = np.polyval(coeffs, x)
        ss_res = np.sum((deviations - predicted) ** 2)
        ss_tot = np.sum((deviations - deviations.mean()) ** 2)
        r_squared = 1 - ss_res / ss_tot if ss_tot > 0 else 0
        r_squared = max(0, min(1, r_squared))

        if r_squared < 0.3:
            return RiskTrend.VOLATILE, 1 - r_squared
        if slope < -0.5:
            return RiskTrend.IMPROVING, r_squared
        elif slope > 0.5:
            return RiskTrend.DETERIORATING, r_squared
        else:
            return RiskTrend.STABLE, r_squared


class ProjectAnalyzer:

    def __init__(self, project_name, project_df, market_df, z_score, api_key=None):
        self.project_name = project_name
        self.raw_project_df = project_df.copy()
        self.market_df = market_df.copy()
        self.z_score = z_score
        self.api_key = api_key

        self.bank = None
        self.df = None
        self.risk_rating = "未评级"
        self.forecast_df = None
        self.stress_results = {}
        self.fig_path = None
        self.use_conservative_mode = False  # 数据质量低时启用保守预测
        self.data_quality_score = 80
        self.risk_trend = None
        self.trend_strength = 0.0
        self.risk_metrics = None  # 风险指标汇总

    def print_sep(self, length=70):
        print(f"{'='*length}")

    def check_data_integrity_enhanced(self) -> ValidationResult:
        """
        多维度检查 + 质量评分 + 详细问题定位
        """
        result = ValidationResult()
        df = self.raw_project_df.copy()
        row_count = len(df)

        result.metadata['总行数'] = row_count
        result.metadata['总列数'] = len(df.columns)
        result.metadata['列名列表'] = list(df.columns)
        required_columns = {
            '日期': 'Date（日期）',
            '当期回收额': 'Actual_Recovery（当期回收金额）',
            '预计收回总金额': 'Total_Recovery（预计收回总金额）',
            '发行银行': 'Bank（发行银行）'
        }

        for col_cn, col_desc in required_columns.items():
            if col_cn not in df.columns:
                result.add_issue(ValidationIssue(
                    severity=Severity.ERROR,
                    category="缺失必填列",
                    message=f"缺少必填列 '{col_cn}' ({col_desc})",
                    suggestion="检查Excel是否包含该列，或修改列名映射配置"
                ))

        if result.has_fatal_error():
            result.quality_score = 0
            result.print_report(self.project_name)
            return result

        # 设置银行名称
        self.bank = df['发行银行'].iloc[0]
        result.metadata['发行银行'] = self.bank

        # 2.1 日期格式检查
        try:
            parsed_dates = pd.to_datetime(df['日期'], errors='coerce')
            invalid_date_mask = parsed_dates.isna()
            invalid_date_rows = df[invalid_date_mask].index.tolist()

            if invalid_date_rows:
                invalid_samples = df.loc[invalid_date_rows[:3], '日期'].tolist()
                result.add_issue(ValidationIssue(
                    severity=Severity.WARNING,
                    category="日期格式异常",
                    message=f"{len(invalid_date_rows)} 行日期无法解析（示例: {invalid_samples}）",
                    affected_rows=invalid_date_rows,
                    suggestion="检查是否为标准日期格式"
                ))
                result.quality_score -= min(len(invalid_date_rows) * 5, 30)  # 最多扣30分
                result.repair_log.append(f"日期解析失败行已设置为NaT，后续将尝试插值处理")

            # 检查时间序列连续性（季度数据）
            valid_dates = parsed_dates.dropna().sort_values()
            if len(valid_dates) >= 2:
                date_diffs = valid_dates.diff().dropna()
                irregular_gaps = date_diffs[abs(date_diffs.dt.days - 90) > 20]
                if len(irregular_gaps) > 0:
                    gap_examples = [
                        f"{valid_dates.iloc[i].strftime('%Y-%m-%d')} ~ {valid_dates.iloc[i+1].strftime('%Y-%m-%d')}"
                        for i in irregular_gaps.index[:3]
                    ]
                    result.add_issue(ValidationIssue(
                        severity=Severity.INFO,
                        category="时间序列",
                        message=f"检测到 {len(irregular_gaps)} 个非标准季度间隔",
                        suggestion=f"间隔示例: {', '.join(gap_examples)}；如确为缺失期数，模型将自动插值"
                    ))

        except Exception as e:
            result.add_issue(ValidationIssue(
                severity=Severity.ERROR,
                category="日期解析失败",
                message=f"日期列处理异常: {str(e)}",
                suggestion="检查日期列是否包含非日期格式的数据"
            ))

        numeric_cols = ['当期回收额', '预计收回总金额']
        for col in numeric_cols:
            non_numeric = pd.to_numeric(df[col], errors='coerce').isna() & df[col].notna()
            if non_numeric.any():
                bad_rows = df[non_numeric].index.tolist()
                result.add_issue(ValidationIssue(
                    severity=Severity.ERROR,
                    category="类型错误",
                    message=f"'{col}' 列包含 {len(bad_rows)} 个非数值数据",
                    affected_rows=bad_rows,
                    suggestion="请检查并修正为纯数字格式"
                ))

        for col in ['当期回收额', '预计收回总金额']:
            if col in df.columns:
                negative_mask = pd.to_numeric(df[col], errors='coerce') < 0
                negative_rows = df[negative_mask].index.tolist()
                if negative_rows:
                    result.add_issue(ValidationIssue(
                        severity=Severity.ERROR,
                        category="业务逻辑错误",
                        message=f"'{col}' 列发现 {len(negative_rows)} 行负数",
                        affected_rows=negative_rows,
                        suggestion="金额不能为负，检查是否为数据录入错误"
                    ))

        zero_total_mask = pd.to_numeric(df['预计收回总金额'], errors='coerce').fillna(0) == 0
        if zero_total_mask.any():
            zero_rows = df[zero_total_mask].index.tolist()
            result.add_issue(ValidationIssue(
                severity=Severity.ERROR,
                category="业务逻辑错误",
                message=f"'预计收回总金额' 列为0或缺失: {len(zero_rows)} 行",
                affected_rows=zero_rows,
                suggestion="预计收回总金额必须大于0"
            ))

        df['Temp_Rate'] = pd.to_numeric(df['当期回收额'], errors='coerce') / \
                          pd.to_numeric(df['预计收回总金额'], errors='coerce').replace(0, np.nan)
        over_100_mask = df['Temp_Rate'] > 1.0
        if over_100_mask.any():
            over_rows = df[over_100_mask].index.tolist()
            over_values = df.loc[over_rows[:3], 'Temp_Rate'].tolist()
            result.add_issue(ValidationIssue(
                severity=Severity.WARNING,
                category="异常值",
                message=f"{len(over_rows)} 行当期回收率超过100%（即单次回收 > 总预期）",
                affected_rows=over_rows,
                suggestion=f"请核实是否为数据错误（示例行回收率: {[f'{v*100:.1f}%' for v in over_values]}），或确为提前全额回收"
            ))
            result.quality_score -= min(len(over_rows) * 3, 15)

        df_sorted = df.sort_values('日期') if '日期' in df.columns else df
        df_sorted['Temp_Cumulative'] = df_sorted['当期回收额'].cumsum()
        df_sorted['Temp_Cumulative_Rate'] = df_sorted['Temp_Cumulative'] / df_sorted['预计收回总金额']
        over_cumulative = df_sorted[df_sorted['Temp_Cumulative_Rate'] > 1.0]
        if not over_cumulative.empty:
            last_over = over_cumulative.iloc[-1]
            result.add_issue(ValidationIssue(
                severity=Severity.INFO,
                category="业务提示",
                message=f"累计回收率已达 {last_over['Temp_Cumulative_Rate']*100:.1f}%，超过100%",
                suggestion="项目可能已提前完成全部回收，后续期数预测将趋于0"
            ))

        macro_cols = ['二手房成交量指数', '二手房成交价指数', '法拍成交量指数',
                      '法拍成交价指数', '国房景气指数']

        missing_ratios = {}
        for col in macro_cols:
            if col in df.columns:
                missing_ratio = df[col].isna().sum() / len(df)
                missing_ratios[col] = missing_ratio
            else:
                missing_ratios[col] = 1.0  # 整列缺失

        avg_missing = sum(missing_ratios.values()) / len(macro_cols)

        if avg_missing > 0:
            missing_details = [f"{col.split('指数')[0]}: {ratio*100:.0f}%"
                               for col, ratio in missing_ratios.items() if ratio > 0]

            if avg_missing > 0.5:  # 缺失超过50%
                result.add_issue(ValidationIssue(
                    severity=Severity.WARNING,
                    category="宏观数据严重缺失",
                    message=f"宏观指标平均缺失率 {avg_missing*100:.1f}%（{', '.join(missing_details)}）",
                    suggestion="缺失部分将用历史均值(100)填充"
                ))
                result.quality_score -= avg_missing * 30
            else:
                result.add_issue(ValidationIssue(
                    severity=Severity.INFO,
                    category="宏观数据部分缺失",
                    message=f"宏观指标平均缺失率 {avg_missing*100:.1f}%（{', '.join(missing_details)}）",
                    suggestion="缺失值将用线性插值或均值填充"
                ))
                result.quality_score -= avg_missing * 15

        completely_missing_macros = [col for col, ratio in missing_ratios.items() if ratio == 1.0]
        if completely_missing_macros:
            result.repair_log.append(f"完全缺失的宏观列将默认填充为100: {', '.join(completely_missing_macros)}")

        if row_count < 3:
            result.add_issue(ValidationIssue(
                severity=Severity.WARNING,
                category="数据量不足",
                message=f"历史数据仅 {row_count} 期",
                suggestion="数据量过少会导致趋势预测(R²)和波动率计算不准确"
            ))
            result.quality_score -= 20
        elif row_count >= 8:
            result.metadata['数据充足性'] = "充足（≥8期）"
        else:
            result.metadata['数据充足性'] = f"一般（{row_count}期）"

        result.quality_score = max(0, min(100, result.quality_score))  # 限制在0-100

        if not result.has_fatal_error():
            result.passed = True

        result.print_report(self.project_name)
        return result

    def check_data_integrity(self) -> bool:
        """
        """
        result = self.check_data_integrity_enhanced()

        self.data_quality_score = result.quality_score

        if result.passed and result.quality_score < 60:
            print(f"数据质量评分仅 {result.quality_score:.1f} 分，将启用保守预测模式")
            self.use_conservative_mode = True
        else:
            self.use_conservative_mode = False

        return result.passed

    def preprocess_and_scale(self):
        col_map = {
            '日期': 'Date', '项目ID': 'Project_ID', '发行银行': 'Bank',
            '当期回收额': 'Actual_Recovery', '预计收回总金额': 'Total_Recovery',
            '二手房成交量指数': 'Housing_Volume_Index', '二手房成交价指数': 'Housing_Price_Index',
            '法拍成交量指数': 'Auction_Volume_Index', '法拍成交价指数': 'Auction_Price_Index',
            '国房景气指数': 'RE_Climate_Index'
        }

        df = preprocess_data(self.raw_project_df, col_map, is_project_data=True)
        benchmarks = calculate_benchmarks_by_lifecycle(self.bank, self.market_df)
        df = pd.merge(df, benchmarks, on='Month_Since_Start', how='left').ffill().bfill()

        df['Baseline_Rate'] = (0.7 * df['Same_Bank_Peer_Rate'] + 0.3 * df['Market_Average_Rate'])
        scaler = StandardScaler()

        macro_cols_original = ['Housing_Volume_Index', 'Housing_Price_Index', 'Auction_Volume_Index',
                               'Auction_Price_Index', 'RE_Climate_Index']
        macro_cols_scaled = ['Volume_Scaled', 'Price_Scaled', 'Auction_Volume_Scaled', 'Auction_Price_Scaled',
                             'RE_Climate_Scaled']
        df[macro_cols_scaled] = scaler.fit_transform(df[macro_cols_original])

        df['Disposal_Environment_Index'] = (
                df['Volume_Scaled'] * MACRO_WEIGHTS['Housing_Volume'] +
                df['Price_Scaled'] * MACRO_WEIGHTS['Housing_Price'] +
                df['Auction_Volume_Scaled'] * MACRO_WEIGHTS['Auction_Volume'] +
                df['Auction_Price_Scaled'] * MACRO_WEIGHTS['Auction_Price'] +
                df['RE_Climate_Scaled'] * MACRO_WEIGHTS['RE_Climate'])

        env_index_mean = df['Disposal_Environment_Index'].mean()
        env_index_std = df['Disposal_Environment_Index'].std()
        df['Market_Multiplier'] = 1 + ((df['Disposal_Environment_Index'] - env_index_mean) / (
            env_index_std if env_index_std > 0 else 1)) * 0.2

        self.df = df

    def calculate_risk_enhanced(self):
        df = self.df.copy()

        if 'Cumulative_Baseline_Rate' not in df.columns:
            df['Cumulative_Baseline_Rate'] = (
                0.7 * df['Same_Bank_Cumulative_Peer_Rate'] +
                0.3 * df['Market_Cumulative_Average_Rate']
            )

        df['Deviation_Rate'] = df['Recovery_Rate'] - df['Baseline_Rate']
        df['Cum_Deviation_Rate'] = df['Cum_Recovery_Rate'] - df['Cumulative_Baseline_Rate']

        vol_estimator = VolatilityEstimator()
        df['Base_Volatility_Rate'] = vol_estimator.estimate(df['Deviation_Rate'])
        df['Cumulative_Volatility'] = vol_estimator.estimate(df['Cum_Deviation_Rate'])

        ci_calculator = ConfidenceIntervalCalculator(CONFIDENCE_LEVEL)

        lower_periodic, upper_periodic = ci_calculator.calculate_bootstrap_ci(
            df['Baseline_Rate'], df['Deviation_Rate']
        )
        df['Lower_Bound_Rate'] = (lower_periodic * df['Market_Multiplier']).clip(lower=0)
        df['Upper_Bound_Rate'] = upper_periodic * df['Market_Multiplier']

        lower_cum, upper_cum = ci_calculator.calculate_bootstrap_ci(
            df['Cumulative_Baseline_Rate'], df['Cum_Deviation_Rate']
        )
        df['Cumulative_Lower_Bound'] = lower_cum.clip(lower=0)
        df['Cumulative_Upper_Bound'] = upper_cum
        df['Anomaly_Positive'] = df['Recovery_Rate'] > df['Upper_Bound_Rate']
        df['Anomaly_Negative'] = df['Recovery_Rate'] < df['Lower_Bound_Rate']
        df['Anomaly'] = df['Anomaly_Positive'] | df['Anomaly_Negative']
        df['Anomaly_Severity'] = 0.0
        positive_mask = df['Anomaly_Positive']
        negative_mask = df['Anomaly_Negative']
        df.loc[positive_mask, 'Anomaly_Severity'] = (
            (df.loc[positive_mask, 'Recovery_Rate'] - df.loc[positive_mask, 'Upper_Bound_Rate'])
            / df.loc[positive_mask, 'Base_Volatility_Rate'].replace(0, 0.01)
        )
        df.loc[negative_mask, 'Anomaly_Severity'] = (
            (df.loc[negative_mask, 'Lower_Bound_Rate'] - df.loc[negative_mask, 'Recovery_Rate'])
            / df.loc[negative_mask, 'Base_Volatility_Rate'].replace(0, 0.01)
        )
        def calculate_cvar(series: pd.Series, alpha = 0.05) -> float:
            var = series.quantile(alpha)
            return series[series <= var].mean() if not series[series <= var].empty else var

        df['CVaR_95'] = calculate_cvar(df['Deviation_Rate'])
        trend, trend_strength = TrendAnalyzer.analyze_deviationtrend(df)
        self.risk_trend = trend
        self.trend_strength = trend_strength
        data_quality = getattr(self, 'data_quality_score', 80)
        thresholds = DynamicThresholds.calculate(df, data_quality)
        last = df.iloc[-1]
        primary_deviation = last['Cumulative_Lower_Bound'] - last['Cum_Recovery_Rate']
        recent_weights = np.exp(np.linspace(-1, 0, min(4, len(df))))
        recent_weights = recent_weights / recent_weights.sum()
        recent_deviations = df['Deviation_Rate'].tail(len(recent_weights))
        weighted_recent_deviation = (recent_deviations * recent_weights).sum()
        composite_deviation = 0.7 * primary_deviation + 0.3 * weighted_recent_deviation
        if trend == RiskTrend.DETERIORATING and trend_strength > 0.5:
            composite_deviation *= 1.2
        elif trend == RiskTrend.IMPROVING and trend_strength > 0.5:
            composite_deviation *= 0.9
        if composite_deviation > thresholds.subprime_to_doubtful:
            self.risk_rating = "可疑类"
            ratinglevel = 3
        elif composite_deviation > thresholds.watch_to_subprime:
            self.risk_rating = "次级类"
            ratinglevel = 2
        elif composite_deviation > thresholds.normal_to_watch:
            self.risk_rating = "关注类"
            ratinglevel = 1
        else:
            self.risk_rating = "正常类"
            ratinglevel = 0
        confidence = self.calculate_rating_confidence(df, ratinglevel, trend_strength, data_quality)

        def calculate_max_drawdown(series: pd.Series) -> float:
            rolling_max = series.expanding().max()
            drawdown = (series - rolling_max) / rolling_max.replace(0, np.nan)
            return drawdown.min()
        self.risk_metrics = RiskMetrics(
            periodic_volatility=last['Base_Volatility_Rate'],
            cumulative_volatility=last['Cumulative_Volatility'],
            tail_risk_95=last['CVaR_95'],
            max_drawdown=calculate_max_drawdown(df['Cum_Recovery_Rate']),
            trend=trend,
            trend_strength=trend_strength,
            rating_confidence=confidence
        )
        self.df = df
        self.print_risk_report(thresholds)

    def calculate_rating_confidence(self, df: pd.DataFrame, ratinglevel: int,
                                     trend_strength: float, data_quality: float) -> float:
        """计算评级置信度"""
        n = len(df)
        sample_factor = min(1.0, n / 8)
        quality_factor = data_quality / 100
        trend_factor = trend_strength

        last = df.iloc[-1]
        deviation = last['Cumulative_Lower_Bound'] - last['Cum_Recovery_Rate']
        thresholds = [1, 5, 10]
        if ratinglevel < len(thresholds):
            distance_to_boundary = abs(deviation - thresholds[ratinglevel])
        else:
            distance_to_boundary = deviation - thresholds[-1]
        boundary_factor = min(1.0, distance_to_boundary / 2)

        confidence = 0.25 * sample_factor + 0.25 * quality_factor + 0.25 * trend_factor + 0.25 * boundary_factor
        return round(confidence, 2)

    def print_risk_report(self, thresholds: 'ProjectAnalyzer.DynamicThresholds'):
        """打印风险报告"""
        print()
        print(f"风险计量报告: {self.project_name}")

        m = self.risk_metrics
        print(f"【波动率估计】")
        print(f"当期波动率 (鲁棒): {m.periodic_volatility:.2f}%")
        print(f"累计波动率: {m.cumulative_volatility:.2f}%")
        print(f"尾部风险 CVaR(95%): {m.tail_risk_95:.2f}%")
        print(f"最大回撤: {m.max_drawdown:.2f}%")
        print(f"【趋势分析】")
        trend_icon = {"改善": "", "稳定": "", "恶化": "", "波动": ""}.get(m.trend.value, "")
        print(f"{trend_icon} 趋势: {m.trend.value} (强度: {m.trend_strength:.0%})")
        print(f"【评级阈值】(动态调整)")
        print(f"正常→关注: {thresholds.normal_to_watch:.2f}%")
        print(f"关注→次级: {thresholds.watch_to_subprime:.2f}%")
        print(f"次级→可疑: {thresholds.subprime_to_doubtful:.2f}%")
        print(f"【综合评级】")
        rating_colors = {"正常类": "", "关注类": "", "次级类": "🟠", "可疑类": ""}
        confidence_bar = "█" * int(m.rating_confidence * 10) + "░" * (10 - int(m.rating_confidence * 10))
        print(f"{rating_colors.get(self.risk_rating, '⚪')} {self.risk_rating}")
        print(f"评级置信度: [{confidence_bar}] {m.rating_confidence:.0%}")
        print()

    def calculate_risk(self):
        use_enhanced = getattr(self, '_use_enhanced_risk', True)
        if use_enhanced:
            self.calculate_risk_enhanced()
        else:
            self.calculate_risk_legacy()

    def calculate_risk_legacy(self):
        df = self.df.copy()
        df['Deviation_Rate'] = df['Recovery_Rate'] - df['Baseline_Rate']
        df['Base_Volatility_Rate'] = df['Deviation_Rate'].shift(1).expanding(min_periods=3).std().ffill().bfill()
        df['Upper_Bound_Rate'] = df['Baseline_Rate'] + df['Market_Multiplier'] * df['Base_Volatility_Rate'] * self.z_score
        df['Lower_Bound_Rate'] = df['Baseline_Rate'] - df['Market_Multiplier'] * df['Base_Volatility_Rate'] * self.z_score
        df.loc[df['Lower_Bound_Rate'] < 0, 'Lower_Bound_Rate'] = 0
        df['Anomaly'] = (df['Recovery_Rate'] > df['Upper_Bound_Rate']) | (df['Recovery_Rate'] < df['Lower_Bound_Rate'])
        df['Cumulative_Baseline_Rate'] = (0.7 * df['Same_Bank_Cumulative_Peer_Rate'] + 0.3 * df['Market_Cumulative_Average_Rate'])
        df['Cum_Deviation_Rate'] = df['Cum_Recovery_Rate'] - df['Cumulative_Baseline_Rate']
        df['Cumulative_Volatility'] = df['Cum_Deviation_Rate'].shift(1).expanding(min_periods=3).std().ffill().bfill()
        df['Cumulative_Upper_Bound'] = df['Cumulative_Baseline_Rate'] + df['Cumulative_Volatility'] * self.z_score
        df['Cumulative_Lower_Bound'] = df['Cumulative_Baseline_Rate'] - df['Cumulative_Volatility'] * self.z_score
        df.loc[df['Cumulative_Lower_Bound'] < 0, 'Cumulative_Lower_Bound'] = 0
        last_period_data = df.iloc[-1]
        deviation = last_period_data['Cumulative_Lower_Bound'] - last_period_data['Cum_Recovery_Rate']
        if deviation > 10:
            self.risk_rating = "可疑类"
        elif deviation > 5:
            self.risk_rating = "次级类"
        elif deviation > 1:
            self.risk_rating = "关注类"
        else:
            self.risk_rating = "正常类"
        self.df = df
        print(f"当前综合风险评级: {self.risk_rating}")

    def run_forecast(self):
        use_enhanced = getattr(self, '_use_enhanced_forecast', True)

        if use_enhanced:
            self.forecast_df = forecast_future_enhanced(
                self.df, self.market_df, self.bank, self.z_score
            )
        else:
            self.forecast_df = forecast_future_hybrid(self.df, self.market_df, self.bank, self.z_score)
    def run_stress_test(self):
        self.stress_results = run_stress_test(self.df, self.market_df, self.bank, self.z_score,
                                              self.risk_rating)
    def generate_artifacts(self):
        self.fig_path = plot_analysis_and_forecast(
            self.df, self.forecast_df, self.project_name, CONFIDENCE_LEVEL,
            risk_metrics=getattr(self, 'risk_metrics', None),
            stress_results=self.stress_results
        )
        if self.api_key:
            briefing = build_report_prompt(
                self.df, self.forecast_df, self.project_name,
                self.bank, self.risk_rating, self.stress_results
            )
            try:
                report_dict = generate_ai_report(self.api_key, briefing)
                if report_dict:
                    export_raw_json(report_dict, self.project_name)
                    export_to_word_from_json(report_dict, self.project_name, self.fig_path, risk_metrics=self.risk_metrics)
                    export_to_pdf_from_json(report_dict, self.project_name, self.fig_path, risk_metrics=self.risk_metrics)
            except Exception as e:
                print(f"经过多次重试，AI生成依然失败。详情: {e}")
    def run_pipeline(self):
        """主调控器：按序执行完整的工作流"""
        print(f"分析生命周期: {self.project_name}")
        if not self.check_data_integrity():
            return False
        try:
            self.preprocess_and_scale()
            self.calculate_risk()
            self.run_forecast()
            self.run_stress_test()
            self.generate_artifacts()
        except Exception as e:
            print(f" 项目 {self.project_name} 发生错误: {e}")
            import traceback
            traceback.print_exc()
            return False
        print(f"完成分析: {self.project_name}")
        return True
def main():
    print("NPL ABS投后管理系统")
    api_key = "sk-5bccb5e2b831492db980d3c4474fe0e3"
    print("如需生成模拟数据，运行实模拟py")
    projects_dict, market_df = load_batch_data()
    if projects_dict is None or market_df is None:
        print("操作取消")
        return
    col_map_market = {
        '日期': 'Date',
        '项目ID': 'Project_ID',
        '发行银行': 'Bank',
        '当期回收额': 'Actual_Recovery',
        '预计收回总金额': 'Total_Recovery'
    }
    market_df_processed = preprocess_data(market_df.copy(), col_map_market)
    z_score = stats.norm.ppf(1 - (1 - CONFIDENCE_LEVEL) / 2)

    total_projects = len(projects_dict)

    for i, (project_name, project_df) in enumerate(projects_dict.items()):
        print(f"\n({i + 1}/{total_projects})")

        analyzer = ProjectAnalyzer(
            project_name=project_name,
            project_df=project_df.copy(),
            market_df=market_df_processed.copy(),
            z_score=z_score,
            api_key=api_key
        )
        analyzer.run_pipeline()

    print("\n")
    print("结束")
if __name__ == "__main__":
    main()
